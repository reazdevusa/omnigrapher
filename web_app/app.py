import re
import streamlit as st
import streamlit.components.v1 as components
import requests
import logging
import os
import json
import io
import uuid
import time
import base64
from datetime import datetime
from typing import Dict, Any
from dotenv import load_dotenv

# Load environment variables
load_dotenv()


# Configure logging
logging.basicConfig(
    level=getattr(logging, os.getenv("LOG_LEVEL", "INFO")),
    format=os.getenv("LOG_FORMAT", "%(asctime)s - %(name)s - %(levelname)s - %(message)s")
)
logger = logging.getLogger(__name__)

# Configuration
API_URL = os.getenv("API_URL", "http://localhost:8001")
API_TIMEOUT = int(os.getenv("API_TIMEOUT", "240"))
MAX_RETRIES = int(os.getenv("MAX_RETRIES", "3"))
RETRY_DELAY = int(os.getenv("RETRY_DELAY", "2"))
HISTORY_FILE = os.path.join(os.path.dirname(__file__), "chat_history.json")
CHATS_DIR = os.path.join(os.path.dirname(__file__), "chats")

_EMAIL_RE = re.compile(r"^[\w\.-]+@[\w\.-]+\.\w+$")


def _chats_dir(username: str) -> str:
    d = os.path.join(CHATS_DIR, username)
    os.makedirs(d, exist_ok=True)
    return d


def _sanitize_answer(answer: str) -> str:
    """Detect and truncate looping answers before saving."""
    if len(answer) < 200:
        return answer
    chunk = answer[:80]
    if answer.count(chunk) > 4:
        first = answer.index(chunk)
        second = answer.index(chunk, first + len(chunk))
        return answer[:second].strip() + "\n\n*[Response truncated — repetition detected]*"
    return answer


def _auto_title(question: str) -> str:
    """Generate a short chat title from the first question."""
    q = question.strip().rstrip("?!.").strip()
    return q[:45] + ("…" if len(q) > 45 else "")


def _load_sessions(username: str) -> list:
    """Return list of session metadata dicts sorted by updated_at desc."""
    d = _chats_dir(username)
    sessions = []
    for fname in os.listdir(d):
        if not fname.endswith(".json"):
            continue
        try:
            with open(os.path.join(d, fname), "r", encoding="utf-8") as f:
                s = json.load(f)
            sessions.append(s)
        except Exception:
            pass
    sessions.sort(key=lambda s: s.get("updated_at", ""), reverse=True)
    return sessions


def _load_session(username: str, session_id: str) -> dict:
    path = os.path.join(_chats_dir(username), f"{session_id}.json")
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return {"id": session_id, "title": "New Chat", "pinned": False,
                "messages": [], "updated_at": datetime.now().isoformat()}


def _save_session(username: str, session: dict) -> None:
    session["updated_at"] = datetime.now().isoformat()
    path = os.path.join(_chats_dir(username), f"{session['id']}.json")
    try:
        with open(path, "w", encoding="utf-8") as f:
            json.dump(session, f, ensure_ascii=False, indent=2)
    except Exception as e:
        logger.warning(f"Could not save session: {e}")


def _delete_session(username: str, session_id: str) -> None:
    path = os.path.join(_chats_dir(username), f"{session_id}.json")
    try:
        os.remove(path)
    except Exception:
        pass


def _new_session_id() -> str:
    return uuid.uuid4().hex[:12]


def _load_history() -> list:
    """Legacy: load flat chat_history.json (migration only)."""
    try:
        if os.path.exists(HISTORY_FILE):
            with open(HISTORY_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
    except Exception:
        pass
    return []


def _save_history(history: list) -> None:
    """Legacy: save flat history (kept for migration path)."""
    try:
        sanitized = [{**e, "answer": _sanitize_answer(e["answer"])} for e in history]
        with open(HISTORY_FILE, "w", encoding="utf-8") as f:
            json.dump(sanitized, f, ensure_ascii=False, indent=2)
    except Exception as e:
        logger.warning(f"Could not save chat history: {e}")

# Page configuration
st.set_page_config(
    page_title="AI Knowledge Base",
    page_icon="🧠",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for better styling
st.markdown("""
<link rel="icon" type="image/svg+xml" href="public/favicon.svg">
<style>
    /* Reduce top whitespace on main content */
    .block-container { padding-top: 0.5rem !important; }
    .main-header {
        font-size: 2.5rem;
        font-weight: bold;
        color: #1f77b4;
        margin-bottom: 1rem;
    }
    /* Increase label text size */
    label[data-testid="stLabel"] {
        font-size: 16px !important;
        font-weight: 500 !important;
    }
    .query-container {
        background-color: #f0f2f6;
        padding: 1.5rem;
        border-radius: 10px;
        margin: 1rem 0;
    }
    .response-container {
        background-color: #e8f4f8;
        padding: 1.5rem;
        border-radius: 10px;
        margin: 1rem 0;
        border-left: 4px solid #1f77b4;
    }
    .info-box {
        background-color: #fff3cd;
        padding: 1rem;
        border-radius: 5px;
        border-left: 4px solid #ffc107;
        margin: 1rem 0;
    }
    /* Style main search button specifically */
    div[data-testid="stButton"] > button[kind="primary"] {
        background-color: #1f77b4 !important;
        color: white !important;
        height: 50px !important;
        font-size: 18px !important;
        font-weight: bold !important;
        padding: 10px 20px !important;
        min-height: 50px !important;
        line-height: 50px !important;
        max-width: 200px !important;
        border: none !important;
    }
    div[data-testid="stButton"] > button[kind="primary"]:hover {
        background-color: #155a8a !important;
        color: white !important;
    }
    /* Style example question buttons as text cards */
    button[kind="secondary"] {
        height: auto !important;
        font-size: 15px !important;
        padding: 12px 16px !important;
        background-color: #f8f9fa !important;
        border: 1px solid #e0e0e0 !important;
        border-radius: 8px !important;
        text-align: left !important;
        width: 100% !important;
        white-space: nowrap !important;
        overflow: hidden !important;
        text-overflow: ellipsis !important;
        justify-content: flex-start !important;
    }
    button[kind="secondary"] > span {
        text-align: left !important;
    }
    button[kind="secondary"]:hover {
        background-color: #e9ecef !important;
        border-color: #d0d7e0 !important;
    }
    /* Style conversation history for better text display */
    [data-testid="stMarkdownContainer"] {
        word-wrap: break-word !important;
        overflow-wrap: break-word !important;
    }
    [data-testid="stMarkdownContainer"] p {
        white-space: pre-wrap !important;
        word-break: break-word !important;
    }
</style>
""", unsafe_allow_html=True)

def set_example_query(query: str):
    """Set example query in session state."""
    logger.info(f"Example query selected: {query[:30]}...")
    st.session_state.query_input = query

@st.cache_data(ttl=10)
def check_backend_connection(api_url: str) -> dict:
    """Check backend and Ollama status. Returns dict with keys: connected, ollama_ok, ollama_message."""
    try:
        logger.info(f"Checking backend connection to: {api_url}")
        response = requests.get(f"{api_url}/", timeout=5)
        if response.status_code == 200:
            data = response.json()
            ollama_ok = data.get("ollama_status") == "ok"
            return {
                "connected": True,
                "ollama_ok": ollama_ok,
                "ollama_message": data.get("ollama_message", ""),
            }
        return {"connected": False, "ollama_ok": False, "ollama_message": ""}
    except requests.exceptions.RequestException as e:
        logger.error(f"Backend connection check failed: {e}")
        return {"connected": False, "ollama_ok": False, "ollama_message": ""}

def query_knowledge_base(query: str, api_url: str) -> Dict[str, Any]:
    """Send query to the backend API (non-streaming fallback)."""
    try:
        logger.info(f"Sending query to API: {query[:50]}...")
        response = requests.post(
            f"{api_url}/api/query",
            json={"query": query},
            timeout=API_TIMEOUT
        )
        response.raise_for_status()
        result = response.json()
        logger.info("API query successful")
        return result
    except requests.exceptions.RequestException as e:
        logger.error(f"API query failed: {e}")
        return {"error": f"Failed to connect to API: {str(e)}"}


def _auth_headers(token: str = None) -> dict:
    """Build Authorization header if token is provided."""
    if token:
        return {"Authorization": f"Bearer {token}"}
    return {}


def login_api(username: str, password: str, api_url: str) -> dict:
    """POST /auth/login — returns {access_token, username, role} or {error}."""
    try:
        r = requests.post(
            f"{api_url}/auth/login",
            json={"username": username, "password": password},
            timeout=10
        )
        if r.status_code == 200:
            return r.json()
        return {"error": r.json().get("detail", "Login failed.")}
    except requests.exceptions.RequestException as e:
        return {"error": f"Cannot reach server: {e}"}


def register_api(username: str, password: str, api_url: str, email: str = "", phone: str = "", display_name: str = "") -> dict:
    """POST /auth/register — returns {success, username} or {error}."""
    try:
        payload = {
            "username": username,
            "email": email.strip(),
            "phone": phone.strip(),
            "password": password,
            "confirm_password": password,
        }
        if display_name.strip():
            payload["display_name"] = display_name.strip()
        r = requests.post(
            f"{api_url}/auth/register",
            json=payload,
            timeout=10
        )
        if r.status_code in (200, 201):
            return r.json()
        return {"error": r.json().get("detail", "Registration failed.")}
    except requests.exceptions.RequestException as e:
        return {"error": f"Cannot reach server: {e}"}


def refresh_token_api(refresh_token: str, api_url: str) -> dict:
    """POST /auth/refresh — returns {access_token, refresh_token, username, role} or {error}."""
    try:
        r = requests.post(
            f"{api_url}/auth/refresh",
            json={"refresh_token": refresh_token},
            timeout=10
        )
        if r.status_code == 200:
            return r.json()
        return {"error": r.json().get("detail", "Refresh failed.")}
    except requests.exceptions.RequestException as e:
        return {"error": f"Cannot reach server: {e}"}


def google_oauth_url_api(api_url: str) -> dict:
    """GET /auth/oauth/google — returns {authorization_url} or {error}."""
    try:
        r = requests.get(f"{api_url}/auth/oauth/google", timeout=10)
        if r.status_code == 200:
            return r.json()
        return {"error": r.json().get("detail", "OAuth not available.")}
    except requests.exceptions.RequestException as e:
        return {"error": f"Cannot reach server: {e}"}


def stream_query(query: str, api_url: str, mode: str = "document", history: list = None, token: str = None):
    """Generator that yields tokens from the SSE streaming endpoint."""
    logger.info(f"Streaming query to API: {query[:50]}...")
    with requests.post(
        f"{api_url}/api/query/stream",
        json={"query": query, "mode": mode, "history": history or []},
        headers=_auth_headers(token),
        stream=True,
        timeout=API_TIMEOUT
    ) as response:
        response.raise_for_status()
        # chunk_size=1 forces immediate byte-level reads from the network
        pending = ""
        for chunk in response.iter_content(chunk_size=1):
            if not chunk:
                continue
            try:
                chunk = chunk.decode("utf-8")
            except Exception:
                continue
            pending += chunk
            while "\n" in pending:
                line, pending = pending.split("\n", 1)
                line = line.strip()
                if not line or not line.startswith("data:"):
                    continue
                data_str = line[len("data:"):].strip()
                if data_str == "[DONE]":
                    return
                try:
                    payload = json.loads(data_str)
                    token = payload.get("token", "")
                    if token:
                        yield token
                except json.JSONDecodeError:
                    continue


def upload_documents(files, api_url: str, token: str = None) -> Dict[str, Any]:
    """Upload documents to the backend API."""
    try:
        logger.info(f"Uploading {len(files)} file(s) to API")
        file_tuples = [("files", (f.name, f.getvalue(), f.type)) for f in files]
        response = requests.post(
            f"{api_url}/api/upload",
            files=file_tuples,
            headers=_auth_headers(token),
            timeout=120
        )
        response.raise_for_status()
        return response.json()
    except requests.exceptions.RequestException as e:
        logger.error(f"Upload failed: {e}")
        return {"error": f"Upload failed: {str(e)}"}


def submit_sync_job_api(api_url: str, token: str) -> Dict[str, Any]:
    """Submit an incremental index sync job to the background queue."""
    try:
        r = requests.post(f"{api_url}/api/jobs/sync-index", headers=_auth_headers(token), timeout=30)
        if r.status_code == 200:
            return r.json()
        return {"error": r.json().get("detail", "Failed to submit sync job")}
    except requests.exceptions.RequestException as e:
        return {"error": str(e)}


def submit_rebuild_job_api(api_url: str, token: str) -> Dict[str, Any]:
    """Submit a full index rebuild job to the background queue."""
    try:
        r = requests.post(f"{api_url}/api/jobs/rebuild-index", headers=_auth_headers(token), timeout=30)
        if r.status_code == 200:
            return r.json()
        return {"error": r.json().get("detail", "Failed to submit rebuild job")}
    except requests.exceptions.RequestException as e:
        return {"error": str(e)}


def get_job_status_api(api_url: str, token: str, job_id: str) -> Dict[str, Any]:
    """Get the status of a background indexing job."""
    try:
        r = requests.get(f"{api_url}/api/jobs/{job_id}", headers=_auth_headers(token), timeout=10)
        if r.status_code == 200:
            return r.json()
        return {"error": r.json().get("detail", "Failed to get job status")}
    except requests.exceptions.RequestException as e:
        return {"error": str(e)}


def list_jobs_api(api_url: str, token: str) -> Dict[str, Any]:
    """List recent background indexing jobs for the current user."""
    try:
        r = requests.get(f"{api_url}/api/jobs", headers=_auth_headers(token), timeout=10)
        if r.status_code == 200:
            return r.json()
        return {"error": r.json().get("detail", "Failed to list jobs")}
    except requests.exceptions.RequestException as e:
        return {"error": str(e)}


def sync_index_api(api_url: str, token: str) -> Dict[str, Any]:
    """Trigger incremental index sync on the backend."""
    try:
        r = requests.post(f"{api_url}/api/sync-index", headers=_auth_headers(token), timeout=120)
        if r.status_code == 200:
            return r.json()
        return {"error": r.json().get("detail", "Sync failed")}
    except requests.exceptions.RequestException as e:
        return {"error": str(e)}


def rebuild_index_api(api_url: str, token: str) -> Dict[str, Any]:
    """Trigger full index rebuild on the backend."""
    try:
        r = requests.post(f"{api_url}/api/rebuild-index", headers=_auth_headers(token), timeout=120)
        if r.status_code == 200:
            return r.json()
        return {"error": r.json().get("detail", "Rebuild failed")}
    except requests.exceptions.RequestException as e:
        return {"error": str(e)}


def submit_feedback(api_url: str, token: str, query: str, response: str, mode: str, rating: str, comment: str = "", session_id: str = "", message_id: str = "") -> Dict[str, Any]:
    """Submit a thumbs up/down rating for an answer."""
    try:
        r = requests.post(
            f"{api_url}/api/feedback",
            json={"query": query, "response": response, "mode": mode, "rating": rating,
                  "comment": comment, "session_id": session_id, "message_id": message_id},
            headers=_auth_headers(token),
            timeout=15
        )
        r.raise_for_status()
        return r.json()
    except requests.exceptions.RequestException as e:
        logger.error(f"Feedback submission failed: {e}")
        return {"error": str(e)}


def delete_document(filename: str, api_url: str, token: str = None) -> Dict[str, Any]:
    """Delete a single document via the backend API."""
    try:
        response = requests.delete(
            f"{api_url}/api/documents/{filename}",
            headers=_auth_headers(token),
            timeout=10
        )
        response.raise_for_status()
        return response.json()
    except requests.exceptions.RequestException as e:
        logger.error(f"Delete failed: {e}")
        return {"error": f"Delete failed: {str(e)}"}


def get_document_chunks_api(api_url: str, token: str, filename: str) -> dict:
    """GET /api/documents/{filename}/chunks — returns indexed chunks."""
    try:
        r = requests.get(f"{api_url}/api/documents/{filename}/chunks",
                         headers=_auth_headers(token), timeout=30)
        if r.status_code == 200:
            return r.json()
        return {"error": r.json().get("detail", "Failed to load chunks.")}
    except requests.exceptions.RequestException as e:
        return {"error": str(e)}


def get_document_content_api(api_url: str, token: str, filename: str) -> dict:
    """GET /api/documents/{filename}/content — returns text content."""
    try:
        r = requests.get(f"{api_url}/api/documents/{filename}/content",
                         headers=_auth_headers(token), timeout=30)
        if r.status_code == 200:
            return r.json()
        return {"error": r.json().get("detail", "Failed to load content.")}
    except requests.exceptions.RequestException as e:
        return {"error": str(e)}


def rename_document_api(api_url: str, token: str, old_name: str, new_name: str) -> dict:
    """PUT /api/documents/{filename}/rename."""
    try:
        r = requests.put(f"{api_url}/api/documents/{old_name}/rename",
                         json={"new_name": new_name},
                         headers=_auth_headers(token), timeout=15)
        if r.status_code == 200:
            return r.json()
        return {"error": r.json().get("detail", "Rename failed.")}
    except requests.exceptions.RequestException as e:
        return {"error": str(e)}


def get_profile_api(api_url: str, token: str) -> dict:
    """GET /auth/profile — returns profile dict or {error}."""
    try:
        r = requests.get(f"{api_url}/auth/profile", headers=_auth_headers(token), timeout=10)
        if r.status_code == 200:
            return r.json()
        return {"error": r.json().get("detail", "Failed to load profile.")}
    except requests.exceptions.RequestException as e:
        return {"error": str(e)}


def update_profile_api(api_url: str, token: str, payload: dict) -> dict:
    """PUT /auth/profile — returns updated profile or {error}."""
    try:
        r = requests.put(f"{api_url}/auth/profile", headers=_auth_headers(token), json=payload, timeout=10)
        if r.status_code == 200:
            return r.json()
        return {"error": r.json().get("detail", "Update failed.")}
    except requests.exceptions.RequestException as e:
        return {"error": str(e)}


def admin_health_status_api(api_url: str, token: str) -> dict:
    try:
        r = requests.get(f"{api_url}/api/admin/health", headers=_auth_headers(token), timeout=10)
        if r.status_code == 200:
            return r.json()
        return {"error": r.json().get("detail", "Failed to load health status")}
    except requests.exceptions.RequestException as e:
        return {"error": str(e)}


def admin_widget_config_api(api_url: str, token: str) -> dict:
    try:
        r = requests.get(f"{api_url}/api/admin/widget-config", headers=_auth_headers(token), timeout=10)
        if r.status_code == 200:
            return r.json()
        return {"error": r.json().get("detail", "Failed to load widget config")}
    except requests.exceptions.RequestException as e:
        return {"error": str(e)}


def admin_list_feedback_api(api_url: str, token: str) -> dict:
    try:
        r = requests.get(f"{api_url}/api/admin/feedback", headers=_auth_headers(token), timeout=10)
        if r.status_code == 200:
            return r.json()
        return {"error": r.json().get("detail", "Failed to load feedback")}
    except requests.exceptions.RequestException as e:
        return {"error": str(e)}


def admin_list_users_api(api_url: str, token: str) -> dict:
    try:
        r = requests.get(f"{api_url}/api/admin/users", headers=_auth_headers(token), timeout=10)
        if r.status_code == 200:
            return r.json()
        return {"error": r.json().get("detail", "Failed to load users.")}
    except requests.exceptions.RequestException as e:
        return {"error": str(e)}


def admin_delete_user_api(api_url: str, token: str, username: str) -> dict:
    try:
        r = requests.delete(f"{api_url}/api/admin/users/{username}", headers=_auth_headers(token), timeout=10)
        if r.status_code == 200:
            return r.json()
        return {"error": r.json().get("detail", "Delete failed.")}
    except requests.exceptions.RequestException as e:
        return {"error": str(e)}


def admin_set_role_api(api_url: str, token: str, username: str, role: str) -> dict:
    try:
        r = requests.put(f"{api_url}/api/admin/users/{username}/role",
                         json={"role": role},
                         headers=_auth_headers(token), timeout=10)
        if r.status_code == 200:
            return r.json()
        return {"error": r.json().get("detail", "Role update failed.")}
    except requests.exceptions.RequestException as e:
        return {"error": str(e)}


@st.cache_data(ttl=30)
def list_documents(api_url: str, token: str = None) -> Dict[str, Any]:
    """List documents in the current user's knowledge base."""
    try:
        response = requests.get(
            f"{api_url}/api/documents",
            headers=_auth_headers(token),
            timeout=10
        )
        response.raise_for_status()
        return response.json()
    except requests.exceptions.RequestException as e:
        logger.error(f"Failed to list documents: {e}")
        return {"documents": [], "count": 0}

# ── Country dial code table ────────────────────────────────────────────────
_DIAL_CODES = [
    ("🇧🇩 Bangladesh", "+880"), ("🇺🇸 United States", "+1"),
    ("🇬🇧 United Kingdom", "+44"), ("🇮🇳 India", "+91"),
    ("🇦🇺 Australia", "+61"), ("🇨🇦 Canada", "+1"),
    ("🇩🇪 Germany", "+49"), ("🇫🇷 France", "+33"),
    ("🇯🇵 Japan", "+81"), ("🇨🇳 China", "+86"),
    ("🇸🇦 Saudi Arabia", "+966"), ("🇦🇪 UAE", "+971"),
    ("🇲🇾 Malaysia", "+60"), ("🇸🇬 Singapore", "+65"),
    ("🇵🇰 Pakistan", "+92"), ("🇮🇩 Indonesia", "+62"),
    ("🇧🇷 Brazil", "+55"), ("🇿🇦 South Africa", "+27"),
    ("🇳🇬 Nigeria", "+234"), ("🇪🇬 Egypt", "+20"),
    ("🇹🇷 Turkey", "+90"), ("🇮🇹 Italy", "+39"),
    ("🇪🇸 Spain", "+34"), ("🇷🇺 Russia", "+7"),
    ("🇰🇷 South Korea", "+82"), ("🇵🇭 Philippines", "+63"),
    ("🇻🇳 Vietnam", "+84"), ("🇹🇭 Thailand", "+66"),
    ("🇳🇿 New Zealand", "+64"), ("🇮🇷 Iran", "+98"),
    ("🇮🇶 Iraq", "+964"), ("🇵🇹 Portugal", "+351"),
    ("🇳🇱 Netherlands", "+31"), ("🇸🇪 Sweden", "+46"),
    ("🇳🇴 Norway", "+47"), ("🇩🇰 Denmark", "+45"),
    ("🇵🇱 Poland", "+48"), ("🇺🇦 Ukraine", "+380"),
    ("🇦🇷 Argentina", "+54"), ("🇲🇽 Mexico", "+52"),
]
# United States pinned first, rest alphabetical
_DIAL_PINNED = [("\U0001f1fa\U0001f1f8 United States", "+1")]
_DIAL_REST = sorted(
    [(n, c) for n, c in _DIAL_CODES if n != "\U0001f1fa\U0001f1f8 United States"],
    key=lambda x: x[0]
)
_DIAL_CODES_SORTED = _DIAL_PINNED + _DIAL_REST
_DIAL_LABELS = [f"{name}  ({code})" for name, code in _DIAL_CODES_SORTED]
_DIAL_MAP = {f"{name}  ({code})": code for name, code in _DIAL_CODES_SORTED}

# Country code → exact dial label (handles US vs CA both being +1)
_CC_TO_LABEL = {
    "BD": "\U0001f1e7\U0001f1e9 Bangladesh  (+880)",
    "US": "\U0001f1fa\U0001f1f8 United States  (+1)",
    "CA": "\U0001f1e8\U0001f1e6 Canada  (+1)",
    "GB": "\U0001f1ec\U0001f1e7 United Kingdom  (+44)",
    "IN": "\U0001f1ee\U0001f1f3 India  (+91)",
    "AU": "\U0001f1e6\U0001f1fa Australia  (+61)",
    "DE": "\U0001f1e9\U0001f1ea Germany  (+49)",
    "FR": "\U0001f1eb\U0001f1f7 France  (+33)",
    "JP": "\U0001f1ef\U0001f1f5 Japan  (+81)",
    "CN": "\U0001f1e8\U0001f1f3 China  (+86)",
    "SA": "\U0001f1f8\U0001f1e6 Saudi Arabia  (+966)",
    "AE": "\U0001f1e6\U0001f1ea UAE  (+971)",
    "MY": "\U0001f1f2\U0001f1fe Malaysia  (+60)",
    "SG": "\U0001f1f8\U0001f1ec Singapore  (+65)",
    "PK": "\U0001f1f5\U0001f1f0 Pakistan  (+92)",
    "ID": "\U0001f1ee\U0001f1e9 Indonesia  (+62)",
    "BR": "\U0001f1e7\U0001f1f7 Brazil  (+55)",
    "ZA": "\U0001f1ff\U0001f1e6 South Africa  (+27)",
    "NG": "\U0001f1f3\U0001f1ec Nigeria  (+234)",
    "EG": "\U0001f1ea\U0001f1ec Egypt  (+20)",
    "TR": "\U0001f1f9\U0001f1f7 Turkey  (+90)",
    "IT": "\U0001f1ee\U0001f1f9 Italy  (+39)",
    "ES": "\U0001f1ea\U0001f1f8 Spain  (+34)",
    "RU": "\U0001f1f7\U0001f1fa Russia  (+7)",
    "KR": "\U0001f1f0\U0001f1f7 South Korea  (+82)",
    "PH": "\U0001f1f5\U0001f1ed Philippines  (+63)",
    "VN": "\U0001f1fb\U0001f1f3 Vietnam  (+84)",
    "TH": "\U0001f1f9\U0001f1ed Thailand  (+66)",
    "NZ": "\U0001f1f3\U0001f1ff New Zealand  (+64)",
    "IR": "\U0001f1ee\U0001f1f7 Iran  (+98)",
    "IQ": "\U0001f1ee\U0001f1f6 Iraq  (+964)",
    "PT": "\U0001f1f5\U0001f1f9 Portugal  (+351)",
    "NL": "\U0001f1f3\U0001f1f1 Netherlands  (+31)",
    "SE": "\U0001f1f8\U0001f1ea Sweden  (+46)",
    "NO": "\U0001f1f3\U0001f1f4 Norway  (+47)",
    "DK": "\U0001f1e9\U0001f1f0 Denmark  (+45)",
    "PL": "\U0001f1f5\U0001f1f1 Poland  (+48)",
    "UA": "\U0001f1fa\U0001f1e6 Ukraine  (+380)",
    "AR": "\U0001f1e6\U0001f1f7 Argentina  (+54)",
    "MX": "\U0001f1f2\U0001f1fd Mexico  (+52)",
}


def _detect_country_dial() -> str:
    """Best-effort IP geolocation — returns exact _DIAL_LABELS entry. Falls back to US."""
    default = "\U0001f1fa\U0001f1f8 United States  (+1)"
    try:
        r = requests.get("https://ipapi.co/json/", timeout=3)
        if r.status_code == 200:
            cc = r.json().get("country_code", "US")
            return _CC_TO_LABEL.get(cc, default)
    except Exception:
        pass
    return default


def _format_phone(dial_code: str, local: str) -> str:
    """Combine dial code + local number into formatted international string."""
    digits = "".join(c for c in local if c.isdigit())
    if not digits:
        return ""
    # Group local digits: first 4, then groups of 3
    if len(digits) <= 4:
        formatted_local = digits
    elif len(digits) <= 7:
        formatted_local = digits[:4] + "-" + digits[4:]
    else:
        formatted_local = digits[:4] + "-" + digits[4:7] + "-" + digits[7:11]
    return f"{dial_code} {formatted_local}".strip()


def _export_txt(question: str, answer: str) -> bytes:
    """Export Q&A as plain text."""
    ts = datetime.now().strftime("%Y-%m-%d %H:%M")
    content = f"AI Knowledge Base Export\n{'='*40}\nDate: {ts}\n\nQuestion:\n{question}\n\nAnswer:\n{answer}\n"
    return content.encode("utf-8")


def _export_md(question: str, answer: str) -> bytes:
    """Export Q&A as Markdown."""
    ts = datetime.now().strftime("%Y-%m-%d %H:%M")
    content = f"# AI Knowledge Base Export\n\n**Date:** {ts}\n\n## Question\n\n{question}\n\n## Answer\n\n{answer}\n"
    return content.encode("utf-8")


def _export_docx(question: str, answer: str) -> bytes:
    """Export Q&A as a Word document."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    ts = datetime.now().strftime("%Y-%m-%d %H:%M")

    title = doc.add_heading("AI Knowledge Base Export", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Date: {ts}").runs[0].italic = True
    doc.add_paragraph("")

    q_heading = doc.add_heading("Question", level=2)
    doc.add_paragraph(question)
    doc.add_paragraph("")

    doc.add_heading("Answer", level=2)
    for line in answer.split("\n"):
        doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _export_pdf(question: str, answer: str) -> bytes:
    """Export Q&A as PDF with full Unicode support (Bengali, Arabic, CJK, etc.)."""
    from fpdf import FPDF

    FONT_REGULAR = r"C:\Windows\Fonts\arial.ttf"
    FONT_BOLD    = r"C:\Windows\Fonts\arialbd.ttf"

    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    pdf.add_font("Arial", "", FONT_REGULAR)
    pdf.add_font("Arial", "B", FONT_BOLD)

    # Title
    pdf.set_font("Arial", "B", 16)
    pdf.cell(0, 10, "AI Knowledge Base Export", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.set_font("Arial", "", 10)
    pdf.cell(0, 8, f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(6)

    # Question
    pdf.set_font("Arial", "B", 12)
    pdf.cell(0, 8, "Question:", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Arial", "", 11)
    pdf.multi_cell(0, 7, question)
    pdf.ln(4)

    # Answer
    pdf.set_font("Arial", "B", 12)
    pdf.cell(0, 8, "Answer:", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Arial", "", 11)
    clean = answer.replace("**", "").replace("*", "").replace("__", "")
    pdf.multi_cell(0, 7, clean)

    return bytes(pdf.output())


def main():
    logger.info("Web app starting")
    
    # Initialize session state
    if 'conversation_history' not in st.session_state:
        st.session_state.conversation_history = []
    if 'chat_history' not in st.session_state:
        st.session_state.chat_history = []  # legacy compat
    if 'api_url' not in st.session_state:
        st.session_state.api_url = API_URL
    if 'query_input' not in st.session_state:
        st.session_state.query_input = ""
    if 'backend_connected' not in st.session_state:
        st.session_state.backend_connected = False
    if 'input_counter' not in st.session_state:
        st.session_state.input_counter = 0
    if 'ollama_ok' not in st.session_state:
        st.session_state.ollama_ok = True
    if 'scroll_to_bottom' not in st.session_state:
        st.session_state.scroll_to_bottom = True
    if 'auth_token' not in st.session_state:
        st.session_state.auth_token = None
    if 'auth_username' not in st.session_state:
        st.session_state.auth_username = None
    if 'auth_role' not in st.session_state:
        st.session_state.auth_role = None
    if 'auth_refresh_token' not in st.session_state:
        st.session_state.auth_refresh_token = None
    if 'current_page' not in st.session_state:
        st.session_state.current_page = "chat"
    if 'active_session_id' not in st.session_state:
        st.session_state.active_session_id = None
    if 'chat_search' not in st.session_state:
        st.session_state.chat_search = ""
    if 'library_items' not in st.session_state:
        _lib_path = os.path.join(CHATS_DIR, "library.json")
        try:
            st.session_state.library_items = json.load(open(_lib_path, encoding="utf-8")) if os.path.exists(_lib_path) else []
        except Exception:
            st.session_state.library_items = []
    if 'projects' not in st.session_state:
        _proj_path = os.path.join(CHATS_DIR, "projects.json")
        try:
            st.session_state.projects = json.load(open(_proj_path, encoding="utf-8")) if os.path.exists(_proj_path) else []
        except Exception:
            st.session_state.projects = []

    # ── Token persistence via st.query_params ─────────────────────────────────
    # On refresh, restore auth state from query params (_t, _u, _r, _rt).
    # Immediately strip any extra params Streamlit may have appended (e.g. _user=user-profile).
    _ALLOWED_PARAMS = {"_t", "_u", "_r", "_rt"}
    params = st.query_params
    saved_token    = params.get("_t")
    saved_user     = params.get("_u")
    saved_role     = params.get("_r")
    saved_refresh  = params.get("_rt")

    # Remove any unexpected params Streamlit injected
    _extra = [k for k in params.keys() if k not in _ALLOWED_PARAMS]
    for k in _extra:
        del st.query_params[k]

    if st.session_state.auth_token is None and saved_token and saved_user:
        st.session_state.auth_token        = saved_token
        st.session_state.auth_refresh_token = saved_refresh
        st.session_state.auth_username     = saved_user
        st.session_state.auth_role          = saved_role or "user"
        # Restore company branding from profile after page reload
        try:
            _rp = get_profile_api(st.session_state.api_url, saved_token)
            st.session_state.company_name    = _rp.get("company_name", "")
            st.session_state.company_logo    = _rp.get("company_logo", "")
            st.session_state.company_website = _rp.get("company_website", "")
            st.session_state.profile_photo   = _rp.get("profile_photo", "")
        except Exception:
            pass

    # ── Auto-refresh token before it expires ─────────────────────────────────
    if st.session_state.auth_token and st.session_state.get("auth_refresh_token"):
        try:
            import base64
            _parts = st.session_state.auth_token.split(".")
            if len(_parts) == 3:
                _payload = json.loads(base64.urlsafe_b64decode(_parts[1] + "==").decode("utf-8"))
                _exp = _payload.get("exp")
                if _exp and _exp - time.time() < 300:  # refresh within 5 min of expiry
                    _refresh_res = refresh_token_api(st.session_state.auth_refresh_token, st.session_state.api_url)
                    if "error" not in _refresh_res:
                        st.session_state.auth_token = _refresh_res["access_token"]
                        st.session_state.auth_refresh_token = _refresh_res.get("refresh_token", "")
                        st.query_params["_t"] = _refresh_res["access_token"]
                    else:
                        # Refresh failed — force re-login
                        st.session_state.auth_token = None
                        st.session_state.auth_refresh_token = None
                        st.session_state.auth_username = None
                        st.session_state.auth_role = None
                        st.query_params.pop("_t", None)
                        st.warning("Session expired. Please log in again.")
                        st.rerun()
        except Exception:
            pass

    # ── Login gate ────────────────────────────────────────────────────────────
    # When auth is enabled (API_KEY_REQUIRED=true on backend), require login.
    # We detect auth mode by checking if /auth/me returns 401 for unauthenticated requests.
    if st.session_state.auth_token is None:
        # Check if backend requires auth
        try:
            probe = requests.get(f"{st.session_state.api_url}/auth/me", timeout=5)
            auth_required = (probe.status_code == 401)
        except Exception:
            auth_required = False

        if auth_required:
            st.markdown("""
            <style>
            /* Collapse top padding, hide sidebar on login screen */
            .block-container { padding-top: 1rem !important; max-width: 480px !important; margin: auto !important; }
            section[data-testid="stSidebar"] { display: none !important; }
            /* Card header banner */
            .auth-banner {
                background: linear-gradient(135deg, #1f77b4 0%, #155a8a 100%);
                border-radius: 12px 12px 0 0;
                padding: 2rem 1.5rem 1.5rem 1.5rem;
                text-align: center;
                margin-bottom: 0;
            }
            .auth-title {
                font-size: 2.4rem;
                font-weight: 800;
                color: white;
                margin: 0 0 0.3rem 0;
                line-height: 1.2;
            }
            .auth-subtitle {
                color: rgba(255,255,255,0.85);
                font-size: 1.05rem;
                font-weight: 400;
                margin: 0;
            }
            /* Card body */
            .auth-body {
                background: white;
                border: 1px solid #e0e0e0;
                border-top: none;
                border-radius: 0 0 12px 12px;
                padding: 1.5rem 1.5rem 1.2rem 1.5rem;
                box-shadow: 0 4px 16px rgba(0,0,0,0.10);
            }
            .auth-body label { font-size: 14px !important; }
            </style>
            """, unsafe_allow_html=True)

            st.markdown("""
            <div class="auth-banner">
                <p class="auth-title">🧠 AI Knowledge Base</p>
                <p class="auth-subtitle">Sign in to your account to continue</p>
            </div>
            <div class="auth-body">
            """, unsafe_allow_html=True)

            login_tab, reg_tab = st.tabs(["Login", "Register"])

            with login_tab:
                lu = st.text_input("Username", key="login_user")
                lp = st.text_input("Password", type="password", key="login_pass")
                if st.button("Login", type="primary", key="login_btn", use_container_width=True):
                    if lu and lp:
                        res = login_api(lu, lp, st.session_state.api_url)
                        if "error" in res:
                            st.error(res["error"])
                        else:
                            st.session_state.auth_token = res["access_token"]
                            st.session_state.auth_refresh_token = res.get("refresh_token", "")
                            st.session_state.auth_username = res["username"]
                            st.session_state.auth_role = res["role"]
                            st.query_params["_t"] = res["access_token"]
                            st.query_params["_u"] = res["username"]
                            st.query_params["_r"] = res["role"]
                            # Pre-load company branding and profile photo
                            _prof = get_profile_api(st.session_state.api_url, res["access_token"])
                            st.session_state.company_name    = _prof.get("company_name", "")
                            st.session_state.company_logo    = _prof.get("company_logo", "")
                            st.session_state.company_website = _prof.get("company_website", "")
                            st.session_state.profile_photo   = _prof.get("profile_photo", "")
                            st.rerun()
                    else:
                        st.warning("Enter username and password.")

                st.markdown("---")
                _oauth_res = google_oauth_url_api(st.session_state.api_url)
                if "error" not in _oauth_res:
                    st.link_button("Sign in with Google", _oauth_res["authorization_url"],
                                   use_container_width=True, key="google_login_btn")
                else:
                    st.caption("Google sign-in is not configured.")

            with reg_tab:
                ru = st.text_input("Username", key="reg_user")
                rem = st.text_input("Email", key="reg_email")

                # Country dial code + local number
                if "reg_detected_dial" not in st.session_state:
                    st.session_state.reg_detected_dial = _detect_country_dial()
                _reg_default_label = st.session_state.reg_detected_dial
                if _reg_default_label not in _DIAL_LABELS:
                    _reg_default_label = _DIAL_LABELS[0]
                _reg_idx = _DIAL_LABELS.index(_reg_default_label)
                rdc1, rdc2 = st.columns([2, 3])
                with rdc1:
                    reg_dial_label = st.selectbox("Country", options=_DIAL_LABELS, index=_reg_idx,
                                                  key="reg_dial_code", label_visibility="collapsed")
                    reg_dial = _DIAL_MAP[reg_dial_label]
                with rdc2:
                    reg_local = st.text_input("Phone number", key="reg_phone_local",
                                              placeholder="e.g. 1712 345 678", label_visibility="collapsed")
                rph = _format_phone(reg_dial, reg_local)
                if rph:
                    st.caption(f"📱 Full number: `{rph}`")

                rdisp = st.text_input("Display Name (optional)", key="reg_display")
                rp = st.text_input("Password (min 8 chars, upper/lower/digit)", type="password", key="reg_pass")
                rp2 = st.text_input("Confirm Password", type="password", key="reg_pass2")
                if st.button("Register", type="primary", key="reg_btn", use_container_width=True):
                    if not (ru and rem and reg_local and rp and rp2):
                        st.warning("Please fill in username, email, phone, password, and confirmation.")
                    elif rp != rp2:
                        st.error("Passwords do not match.")
                    elif len(rp) < 8 or not any(c.isupper() for c in rp) or not any(c.islower() for c in rp) or not any(c.isdigit() for c in rp):
                        st.error("Password must be at least 8 characters and contain uppercase, lowercase, and a digit.")
                    elif not _EMAIL_RE.match(rem.strip()):
                        st.error("Please enter a valid email address.")
                    elif not rph:
                        st.error("Please enter a valid phone number.")
                    else:
                        res = register_api(ru, rp, st.session_state.api_url, email=rem.strip(), phone=rph, display_name=rdisp)
                        if "error" in res:
                            st.error(res["error"])
                        else:
                            st.success(f"✅ Account created! Switch to Login tab to sign in as **{ru}**.")

            st.markdown("</div>", unsafe_allow_html=True)
            return  # stop rendering the main app until logged in

    # Only check connection on first load — not on every radio/button interaction
    if 'connection_checked' not in st.session_state:
        status = check_backend_connection(st.session_state.api_url)
        st.session_state.backend_connected = status["connected"]
        st.session_state.ollama_ok = status.get("ollama_ok", True)
        st.session_state.connection_checked = True
        st.session_state.connection_status = status
    else:
        status = st.session_state.get("connection_status", {"connected": st.session_state.backend_connected, "ollama_ok": st.session_state.ollama_ok})

    # Auto-scroll is handled inside the chat fragment; the main app must not
    # consume the scroll flag before the fragment sees it.

    # ── Header with company branding ──────────────────────────────────────
    _co_name    = st.session_state.get("company_name", "")
    _co_logo    = st.session_state.get("company_logo", "")
    _co_website = st.session_state.get("company_website", "")
    if _co_name or _co_logo or _co_website:
        _logo_html = (f'<img src="{_co_logo}" style="height:140px;object-fit:contain;display:block;margin:0 auto 2px auto" />'
                      if _co_logo else "")
        _name_html = (f'<div style="font-size:2.25rem;font-weight:800;color:#1a1a1a;line-height:1.15;margin:0;">{_co_name}</div>'
                      if _co_name else "")
        _site_html  = (f'<a href="{_co_website}" target="_blank" style="font-size:0.95rem;color:#1f77b4;text-decoration:none;">🌐 {_co_website}</a>'
                       if _co_website else "")
        st.markdown(f"""
        <div style="text-align:center;margin-top:-0.8rem;padding:0 0 0.2rem 0;line-height:1.3;">
          {_logo_html}
          {_name_html}
          <div style="font-size:1.3rem;font-weight:700;color:#555;margin:1px 0;">🧠 AI Knowledge Base</div>
          {('<div style="margin:1px 0;">' + _site_html + '</div>') if _site_html else ''}
        </div>""", unsafe_allow_html=True)
    else:
        st.markdown('<h1 class="main-header">🧠 AI Knowledge Base</h1>', unsafe_allow_html=True)

    # ── Ollama warning banner (shown above everything else) ──────────────────
    if status["connected"] and not status["ollama_ok"]:
        st.error(
            "⚠️ **Ollama is not running.**\n\n"
            "The AI engine (Ollama) needs to be started before you can ask questions.\n\n"
            "**What to do:**\n"
            "1. Find the **Ollama** icon in your system tray (bottom-right of taskbar) and click it, OR\n"
            "2. Open a terminal and type: `ollama serve`\n\n"
            "Once Ollama is running, refresh this page and you're good to go! 🚀"
        )
    # ────────────────────────────────────────────────────────────────────────

    # ── Sidebar CSS ──────────────────────────────────────────────────────────
    st.markdown("""
    <style>
    /* ── Disable rerun dimming overlay ── */
    [data-testid="stAppViewBlockContainer"] { opacity: 1 !important; }
    .stApp > div[style*="opacity"] { opacity: 1 !important; }
    div[data-testid="stStatusWidget"] { display: none !important; }
    iframe[title="streamlit_analytics"] { display: none !important; }
    .stSpinner { display: none !important; }
    /* The actual gray overlay Streamlit injects on rerun */
    .stApp [class*="overlayText"] { display: none !important; }
    body > div.stApp::after { display: none !important; }
    /* Hide the running indicator dots */
    [data-testid="stDecoration"] { display: none !important; }
    /* Prevent opacity fade on main content during rerun */
    .main .block-container { opacity: 1 !important; transition: none !important; }
    section[data-testid="stSidebar"] { opacity: 1 !important; transition: none !important; }

    /* ── Sidebar width ── */
    section[data-testid="stSidebar"] { min-width:260px !important; max-width:300px !important; }
    section[data-testid="stSidebar"] > div:first-child { padding-top:0.5rem !important; }

    /* ── Collapse ALL gaps between widgets in sidebar ── */
    section[data-testid="stSidebar"] [data-testid="stVerticalBlock"] { gap:2px !important; }
    section[data-testid="stSidebar"] .stButton > button {
        margin-top: 0 !important; margin-bottom: 0 !important;
        padding-top: 4px !important; padding-bottom: 4px !important;
    }
    section[data-testid="stSidebar"] hr { margin: 4px 0 !important; }
    section[data-testid="stSidebar"] [data-testid="stCaptionContainer"] {
        padding: 2px 4px 0px 4px !important; margin: 0 !important;
    }

    /* ── New Chat button (key=btn_new_chat) ── */
    section[data-testid="stSidebar"] [data-testid="stBaseButton-secondary"][key="btn_new_chat"],
    button[kind="secondary"][data-testid="stBaseButton-secondary"] { }
    section[data-testid="stSidebar"] button[kind="secondary"]:first-of-type {
        background: #1f77b4 !important; color: white !important;
        border-radius: 8px !important; font-weight: 600 !important;
        border: none !important;
    }

    /* ── All sidebar buttons: compact, left-aligned, no border ── */
    section[data-testid="stSidebar"] .stButton > button {
        text-align: left !important;
        border-radius: 6px !important;
        font-size: 13px !important;
        line-height: 1.3 !important;
        width: 100% !important;
        border: 1px solid transparent !important;
    }

    /* ── Active nav page highlight ── */
    section[data-testid="stSidebar"] button[data-active="true"] {
        background: #e8f0fe !important; color: #1f77b4 !important; font-weight:700 !important;
    }

    /* ── Avatar strip ── */
    .sb-avatar {
        display: flex; align-items: center; gap: 8px;
        padding: 5px 4px; border-top: 1px solid #e5e7eb; margin-top: 3px;
    }
    .sb-avatar-circle {
        width: 32px; height: 32px; border-radius: 50%;
        background: #1f77b4; color: white; font-weight: 700; font-size: 13px;
        display: flex; align-items: center; justify-content: center; flex-shrink: 0;
    }
    .sb-avatar-name { font-size: 12px; font-weight: 600; color: #111; }
    .sb-avatar-role { font-size: 10px; color: #888; }
    </style>
    """, unsafe_allow_html=True)

    with st.sidebar:
        # ── New Chat button ──────────────────────────────────────────────────
        if st.button("✏️  New Chat", use_container_width=True, key="btn_new_chat"):
            new_id = _new_session_id()
            st.session_state.active_session_id = new_id
            st.session_state.current_page = "chat"
            st.rerun()

        # ── Nav items: Library / Projects / More ──────────────────────────────
        _nav = [
            ("library",  "📚 Library"),
            ("projects", "🗂️ Projects"),
            ("more",     "⋯  More"),
        ]
        for _page_id, _page_label in _nav:
            if st.button(_page_label, key=f"nav_{_page_id}", use_container_width=True):
                st.session_state.current_page = _page_id
                st.rerun()

        # ── Search ───────────────────────────────────────────────────────────
        search_q = st.text_input("🔍 Search chats", value="",
                                  placeholder="Search chats…", key="sidebar_search",
                                  label_visibility="collapsed")

        # ── Chat session list ────────────────────────────────────────────────
        _uname = st.session_state.auth_username or "guest"
        all_sessions = _load_sessions(_uname)

        # Migrate legacy flat history on first login
        if not all_sessions:
            legacy = _load_history()
            if legacy:
                sid = _new_session_id()
                title = _auto_title(legacy[0]["question"]) if legacy else "Previous Chat"
                _save_session(_uname, {
                    "id": sid, "title": title, "pinned": False,
                    "messages": legacy, "updated_at": datetime.now().isoformat()
                })
                if st.session_state.active_session_id is None:
                    st.session_state.active_session_id = sid
                all_sessions = _load_sessions(_uname)

        # Auto-select most recent session on first load
        if st.session_state.active_session_id is None and all_sessions:
            st.session_state.active_session_id = all_sessions[0]["id"]

        # Filter by search
        if search_q.strip():
            sq = search_q.lower()
            all_sessions = [s for s in all_sessions
                            if sq in s.get("title", "").lower()
                            or any(sq in m.get("question", "").lower()
                                   for m in s.get("messages", []))]

        pinned   = [s for s in all_sessions if s.get("pinned")]
        recents  = [s for s in all_sessions if not s.get("pinned")]

        def _session_button(s):
            is_active = s["id"] == st.session_state.active_session_id
            title = s.get("title", "New Chat")
            label = ("📌 " if s.get("pinned") else "") + title
            _rename_key = f"renaming_{s['id']}"
            _opts_key   = f"show_opts_{s['id']}"

            # ── If rename mode is active, show inline text input ─────────────
            if st.session_state.get(_rename_key, False):
                new_title = st.text_input(
                    "Rename", value=title,
                    key=f"rename_input_{s['id']}",
                    label_visibility="collapsed"
                )
                rc1, rc2 = st.columns(2)
                with rc1:
                    if st.button("✅ Save", key=f"rename_save_{s['id']}", use_container_width=True):
                        if new_title.strip():
                            s["title"] = new_title.strip()
                            _save_session(_uname, s)
                        st.session_state[_rename_key] = False
                        st.session_state[_opts_key] = False
                        st.rerun()
                with rc2:
                    if st.button("✖ Cancel", key=f"rename_cancel_{s['id']}", use_container_width=True):
                        st.session_state[_rename_key] = False
                        st.rerun()
                return  # don't render the normal row while editing

            c1, c2 = st.columns([5, 1])
            with c1:
                if st.button(label, key=f"sess_{s['id']}", use_container_width=True):
                    st.session_state.active_session_id = s["id"]
                    st.session_state.current_page = "chat"
                    st.rerun()
            with c2:
                if st.button("⋯", key=f"opt_{s['id']}", help="Rename / Pin / Delete"):
                    st.session_state[_opts_key] = not st.session_state.get(_opts_key, False)

            if st.session_state.get(_opts_key, False):
                pin_label = "📌 Unpin" if s.get("pinned") else "📌 Pin"
                oc1, oc2, oc3 = st.columns(3)
                with oc1:
                    if st.button("✏️ Rename", key=f"rename_{s['id']}", use_container_width=True):
                        st.session_state[_rename_key] = True
                        st.session_state[_opts_key] = False
                        st.rerun()
                with oc2:
                    if st.button(pin_label, key=f"pin_{s['id']}", use_container_width=True):
                        s["pinned"] = not s.get("pinned", False)
                        _save_session(_uname, s)
                        st.session_state[_opts_key] = False
                        st.rerun()
                with oc3:
                    if st.button("🗑️", key=f"del_s_{s['id']}", use_container_width=True, help="Delete"):
                        _delete_session(_uname, s["id"])
                        if st.session_state.active_session_id == s["id"]:
                            st.session_state.active_session_id = None
                        st.session_state[_opts_key] = False
                        st.rerun()

        if pinned:
            st.caption("📌 PINNED")
            for s in pinned:
                _session_button(s)

        if recents:
            st.caption("RECENTS")
            for s in recents[:30]:
                _session_button(s)

        if not pinned and not recents:
            st.caption("No chats yet. Start a new one!")

        st.markdown("---")

        # ── Upload Documents ─────────────────────────────────────────────────
        with st.expander("📂 Upload Documents", expanded=False):
            uploaded_files = st.file_uploader(
                "Upload files to the knowledge base",
                type=["pdf", "txt", "docx", "doc", "md", "csv", "xlsx", "pptx", "html", "epub"],
                accept_multiple_files=True,
                help="Supported: PDF, TXT, DOCX, DOC, MD, CSV, XLSX, PPTX, HTML, EPUB"
            )
            if st.button("⬆️ Upload to Knowledge Base", use_container_width=True,
                         disabled=(not st.session_state.backend_connected or not uploaded_files)):
                with st.spinner(f"Uploading {len(uploaded_files)} file(s)..."):
                    result = upload_documents(uploaded_files, st.session_state.api_url, token=st.session_state.auth_token)
                if "error" in result:
                    st.error(result["error"])
                else:
                    st.success(result.get("message", "Upload successful!"))
                    if result.get("skipped"):
                        st.warning(f"Skipped: {', '.join(result['skipped'])}")

            # ── Index sync controls ───────────────────────────────────────────
            if st.session_state.backend_connected:
                s1, s2 = st.columns(2)
                with s1:
                    if st.button("🔄 Sync Index", use_container_width=True, help="Index only changed/new files"):
                        with st.spinner("Submitting sync job..."):
                            res = submit_sync_job_api(st.session_state.api_url, st.session_state.auth_token)
                        if "error" in res:
                            st.error(res["error"])
                        else:
                            st.session_state["last_job_id"] = res.get("job_id")
                            st.success("Sync job submitted. Check status below.")
                            st.rerun()
                with s2:
                    if st.button("🔃 Full Rebuild", use_container_width=True, help="Delete and re-index all files"):
                        with st.spinner("Submitting rebuild job..."):
                            res = submit_rebuild_job_api(st.session_state.api_url, st.session_state.auth_token)
                        if "error" in res:
                            st.error(res["error"])
                        else:
                            st.session_state["last_job_id"] = res.get("job_id")
                            st.success("Rebuild job submitted. Check status below.")
                            st.rerun()

                # ── Job status panel ──────────────────────────────────────────
                with st.expander("🛠️ Background Jobs", expanded=True):
                    jobs_res = list_jobs_api(st.session_state.api_url, st.session_state.auth_token)
                    if "error" in jobs_res:
                        st.caption(f"Could not load jobs: {jobs_res['error']}")
                    else:
                        jobs = jobs_res.get("jobs", [])
                        if not jobs:
                            st.caption("No background jobs yet.")
                        else:
                            for job in jobs[:5]:
                                job_status = job.get("status", "pending")
                                status_emoji = {
                                    "pending": "⏳",
                                    "running": "🔄",
                                    "completed": "✅",
                                    "failed": "❌",
                                    "cancelled": "🚫",
                                }.get(job_status, "⚪")
                                job_type = job.get("job_type", "index")
                                created = job.get("created_at", "")[:19].replace("T", " ")
                                result = job.get("result") or {}
                                detail = ""
                                if job_status == "completed":
                                    if "indexed" in result:
                                        detail = f" — indexed {result.get('indexed', 0)}, skipped {result.get('skipped', 0)}, removed {result.get('removed', 0)}"
                                elif job_status == "failed":
                                    detail = f" — {job.get('error_message', '')[:80]}"
                                st.caption(f"{status_emoji} {job_type} | {job_status} | {created}{detail}")

            if st.session_state.backend_connected:
                doc_info = list_documents(st.session_state.api_url, st.session_state.auth_token)
                count = doc_info.get("count", 0)
                st.caption(f"📄 {count} document(s) in knowledge base")
                if count > 0:
                    for _doc in doc_info.get("documents", []):
                        doc = _doc["filename"]
                        doc_status = _doc.get("status", "pending")
                        doc_chunks = _doc.get("chunks", 0)
                        doc_error = _doc.get("error", "")
                        _doc_menu_key   = f"docmenu_{doc}"
                        _doc_rename_key = f"docrename_{doc}"
                        _doc_view_key   = f"docview_{doc}"

                        _status_badge = {
                            "indexed": "🟢 indexed",
                            "indexing": "🟡 indexing",
                            "failed": "🔴 failed",
                            "pending": "⚪ pending",
                        }.get(doc_status, f"⚪ {doc_status}")

                        # ── Name row: clickable label + three-dot ──────────
                        dn1, dn2 = st.columns([5, 1])
                        with dn1:
                            _btn_label = f"📄 {doc}\n\n{_status_badge}  ·  {doc_chunks} chunk(s)"
                            if st.button(_btn_label, key=f"docbtn_{doc}",
                                         use_container_width=True, help="Click to preview"):
                                st.session_state[_doc_view_key] = not st.session_state.get(_doc_view_key, False)
                                st.session_state[_doc_menu_key] = False
                        with dn2:
                            if st.button("⋯", key=f"docdot_{doc}", help="Rename / Delete"):
                                st.session_state[_doc_menu_key] = not st.session_state.get(_doc_menu_key, False)
                                st.session_state[_doc_view_key] = False
                                st.session_state[_doc_rename_key] = False

                        if doc_error and doc_status == "failed":
                            st.error(f"Indexing error for {doc}: {doc_error}")

                        # ── Three-dot action row ───────────────────────────
                        if st.session_state.get(_doc_menu_key, False):
                            dm1, dm2, dm3, dm4 = st.columns([1, 1, 1, 1])
                            with dm1:
                                if st.button("✏️", key=f"docren_{doc}", help="Rename",
                                             use_container_width=True):
                                    st.session_state[_doc_rename_key] = True
                                    st.session_state[_doc_menu_key] = False
                                    st.rerun()
                            with dm2:
                                if st.button("🧩", key=f"docchunks_{doc}", help="View chunks",
                                             use_container_width=True):
                                    st.session_state["doc_chunks_name"] = doc
                                    st.session_state["doc_chunks_data"] = None
                                    st.session_state[_doc_menu_key] = False
                                    st.session_state[_doc_view_key] = False
                                    st.rerun()
                            with dm3:
                                if st.button("🗑️", key=f"docdel_{doc}", help="Delete",
                                             use_container_width=True):
                                    st.session_state._pending_delete_doc = doc
                                    st.session_state[_doc_menu_key] = False
                                    st.rerun()

                        # ── Deletion confirmation dialog ──────────────────
                        if st.session_state.get("_pending_delete_doc") == doc:
                            @st.dialog("Confirm deletion")
                            def _confirm_doc_delete():
                                st.write(f"Are you sure you want to permanently delete **{doc}**?")
                                _c1, _c2 = st.columns(2)
                                with _c1:
                                    if st.button("Yes, delete", type="primary", key="confirm_doc_delete"):
                                        with st.spinner(f"Deleting {doc}..."):
                                            result = delete_document(doc, st.session_state.api_url,
                                                                     token=st.session_state.auth_token)
                                        if "error" in result:
                                            st.error(result["error"])
                                        else:
                                            st.session_state.pop("_pending_delete_doc", None)
                                            st.success(f"'{doc}' deleted.")
                                            list_documents.clear()
                                            st.rerun()
                                with _c2:
                                    if st.button("Cancel", key="cancel_doc_delete"):
                                        st.session_state.pop("_pending_delete_doc", None)
                                        st.rerun()
                            _confirm_doc_delete()

                        # ── Inline rename ──────────────────────────────────
                        if st.session_state.get(_doc_rename_key, False):
                            _new_name = st.text_input(
                                "New name", value=doc,
                                key=f"docnewinput_{doc}",
                                label_visibility="collapsed"
                            )
                            rc1, rc2 = st.columns(2)
                            with rc1:
                                if st.button("💾 Save", key=f"docrenameSave_{doc}",
                                             use_container_width=True):
                                    if _new_name.strip() and _new_name.strip() != doc:
                                        res = rename_document_api(
                                            st.session_state.api_url,
                                            st.session_state.auth_token,
                                            doc, _new_name.strip()
                                        )
                                        if "error" in res:
                                            st.error(res["error"])
                                        else:
                                            st.success(f"Renamed to '{_new_name.strip()}'")
                                            list_documents.clear()
                                            st.session_state[_doc_rename_key] = False
                                            st.rerun()
                                    else:
                                        st.session_state[_doc_rename_key] = False
                                        st.rerun()
                            with rc2:
                                if st.button("✖ Cancel", key=f"docrenameCancel_{doc}",
                                             use_container_width=True):
                                    st.session_state[_doc_rename_key] = False
                                    st.rerun()

                        # ── Document preview (shown in main area via session flag) ──
                        if st.session_state.get(_doc_view_key, False):
                            with st.spinner(f"Loading {doc}..."):
                                _content_res = get_document_content_api(
                                    st.session_state.api_url,
                                    st.session_state.auth_token, doc
                                )
                            if "error" in _content_res:
                                st.error(_content_res["error"])
                            else:
                                # Store for main area rendering
                                st.session_state["doc_preview_name"]    = doc
                                st.session_state["doc_preview_content"] = _content_res.get("content", "")
                                st.session_state["doc_preview_type"]    = _content_res.get("type", "text")
                                st.session_state["doc_preview_pages"]   = _content_res.get("pages", "")

        # ── Connection status (collapsed) ────────────────────────────────────
        with st.expander("⚙️ Settings", expanded=False):
            if status["connected"]:
                if status["ollama_ok"]:
                    st.success("✅ Backend Connected")
                else:
                    st.warning("⚠️ Backend up — Ollama not running")
            else:
                st.error("❌ Backend Disconnected")
            api_url_input = st.text_input("API URL", value=st.session_state.api_url, key="api_url_input")
            st.session_state.api_url = api_url_input

        # ── User avatar at bottom ────────────────────────────────────────────
        if st.session_state.auth_username:
            uname = st.session_state.auth_username
            role  = st.session_state.auth_role or "user"
            initials = "".join(p[0].upper() for p in uname.split()[:2]) or uname[:2].upper()
            _sb_photo = st.session_state.get("profile_photo", "")
            if _sb_photo:
                _avatar_circle = f"""<img src="{_sb_photo}" style="width:32px;height:32px;object-fit:cover;border-radius:50%;" />"""
            else:
                _avatar_circle = f"""<div class="sb-avatar-circle">{initials}</div>"""
            st.markdown(f"""
            <div class="sb-avatar">
              {_avatar_circle}
              <div>
                <div class="sb-avatar-name">{uname}</div>
                <div class="sb-avatar-role">{role.capitalize()}</div>
              </div>
            </div>""", unsafe_allow_html=True)
            if st.session_state.get("auth_role") == "admin":
                if st.button("⚙️ Admin Panel", use_container_width=True, key="btn_admin"):
                    st.session_state.current_page = "admin"
                    st.rerun()
            av1, av2 = st.columns(2)
            with av1:
                if st.button("👤 Profile", use_container_width=True, key="btn_profile"):
                    st.session_state.current_page = "profile"
                    st.rerun()
            with av2:
                if st.button("🚪 Logout", use_container_width=True, key="btn_logout"):
                    st.query_params.clear()
                    for k in ["auth_token", "auth_username", "auth_role", "auth_refresh_token",
                               "active_session_id", "detected_dial", "connection_checked"]:
                        st.session_state.pop(k, None)
                    st.session_state.current_page = "chat"
                    st.rerun()
    
        
    # ── Library page ──────────────────────────────────────────────────────────
    if st.session_state.current_page == "library":
        _lib_path = os.path.join(CHATS_DIR, "library.json")

        def _save_library():
            os.makedirs(CHATS_DIR, exist_ok=True)
            with open(_lib_path, "w", encoding="utf-8") as _f:
                json.dump(st.session_state.library_items, _f, ensure_ascii=False, indent=2)

        # Always reload from disk so list is never stale
        try:
            st.session_state.library_items = json.load(open(_lib_path, encoding="utf-8")) if os.path.exists(_lib_path) else []
        except Exception:
            st.session_state.library_items = []

        st.markdown("## 📚 Library")
        st.caption("Save your favourite AI responses and documents for quick reference.")
        st.markdown("---")

        # ── Add to library — pick from ANY chat session ────────────────────
        _uname_lib = st.session_state.auth_username or "guest"
        _all_lib_sessions = _load_sessions(_uname_lib)  # all chats for this user

        with st.expander("➕ Save a response to Library", expanded=False):
            if _all_lib_sessions:
                # Step 1: pick a chat
                _sess_labels = [s.get("title", "New Chat") for s in _all_lib_sessions]
                _sel_sess_idx = st.selectbox(
                    "Pick a chat", range(len(_sess_labels)),
                    format_func=lambda i: _sess_labels[i],
                    key="lib_pick_sess"
                )
                _chosen_sess = _load_session(_uname_lib, _all_lib_sessions[_sel_sess_idx]["id"])
                _lib_msgs = _chosen_sess.get("messages", [])

                # Step 2: pick a message from that chat
                if _lib_msgs:
                    _choices = [
                        f"Q{i+1}: {m['question'][:70]}…" if len(m['question']) > 70
                        else f"Q{i+1}: {m['question']}"
                        for i, m in enumerate(_lib_msgs)
                    ]
                    _sel_idx = st.selectbox(
                        "Pick a response", range(len(_choices)),
                        format_func=lambda i: _choices[i], key="lib_pick_msg"
                    )
                    _lib_note = st.text_input(
                        "Note / tag (optional)", key="lib_note",
                        placeholder="e.g. Resume template, Rust notes…"
                    )
                    if st.button("📥 Save to Library", key="lib_save_btn", type="primary"):
                        _entry = _lib_msgs[_sel_idx]
                        st.session_state.library_items.append({
                            "id": _new_session_id(),
                            "title": _lib_note.strip() or _auto_title(_entry["question"]),
                            "question": _entry["question"],
                            "answer": _entry["answer"],
                            "chat": _sess_labels[_sel_sess_idx],
                            "saved_at": datetime.now().strftime("%Y-%m-%d %H:%M"),
                        })
                        _save_library()
                        st.success("Saved to Library!")
                        st.rerun()
                else:
                    st.info("This chat has no messages yet.")
            else:
                st.info("No chats found. Start a conversation first, then come back to save responses.")

        st.markdown("---")

        # ── Filter tabs ───────────────────────────────────────────────────
        _lib_tab_all, _lib_tab_search = st.tabs(["All Items", "🔍 Search"])
        with _lib_tab_search:
            _lib_sq = st.text_input("Search library", placeholder="keyword…", key="lib_search",
                                    label_visibility="collapsed")
        _lib_items = st.session_state.library_items
        with _lib_tab_search:
            if _lib_sq.strip():
                _lib_items = [x for x in _lib_items
                              if _lib_sq.lower() in x.get("title","").lower()
                              or _lib_sq.lower() in x.get("question","").lower()
                              or _lib_sq.lower() in x.get("answer","").lower()]

        # ── Display items ─────────────────────────────────────────────────
        def _render_lib_items(items):
            if not items:
                st.markdown("""
                <div style='text-align:center;padding:3rem 1rem;color:#aaa'>
                  <div style='font-size:3rem'>📭</div>
                  <div style='font-size:1.1rem;margin-top:.5rem'>No items saved yet</div>
                  <div style='font-size:.9rem'>Save AI responses from any chat above</div>
                </div>""", unsafe_allow_html=True)
                return
            for item in items:
                with st.container(border=True):
                    hc1, hc2 = st.columns([5, 1])
                    with hc1:
                        st.markdown(f"**{item['title']}**")
                        _src = f"💬 {item['chat']}  ·  " if item.get("chat") else ""
                        st.caption(f"{_src}🕐 {item.get('saved_at', '')}")
                    with hc2:
                        if st.button("🗑️", key=f"lib_del_{item['id']}", help="Remove"):
                            st.session_state.library_items = [
                                x for x in st.session_state.library_items if x["id"] != item["id"]]
                            _save_library()
                            st.rerun()
                    with st.expander("View response"):
                        st.markdown(f"**Q:** {item['question']}")
                        st.markdown("---")
                        st.markdown(item["answer"])
                        ec1, ec2, ec3, ec4 = st.columns(4)
                        _fn = item["title"][:30].replace(" ", "_")
                        with ec1:
                            st.download_button("⬇ TXT", data=_export_txt(item["question"], item["answer"]),
                                file_name=f"{_fn}.txt", mime="text/plain", key=f"lib_txt_{item['id']}")
                        with ec2:
                            st.download_button("⬇ MD", data=_export_md(item["question"], item["answer"]),
                                file_name=f"{_fn}.md", mime="text/markdown", key=f"lib_md_{item['id']}")
                        with ec3:
                            st.download_button("⬇ DOCX", data=_export_docx(item["question"], item["answer"]),
                                file_name=f"{_fn}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key=f"lib_docx_{item['id']}")
                        with ec4:
                            st.download_button("⬇ PDF", data=_export_pdf(item["question"], item["answer"]),
                                file_name=f"{_fn}.pdf", mime="application/pdf", key=f"lib_pdf_{item['id']}")

        with _lib_tab_all:
            _render_lib_items(st.session_state.library_items)
        with _lib_tab_search:
            if _lib_sq.strip():
                _render_lib_items(_lib_items)
        return

    # ── Projects page ──────────────────────────────────────────────────────────
    if st.session_state.current_page == "projects":
        _proj_path = os.path.join(CHATS_DIR, "projects.json")

        def _save_projects():
            os.makedirs(CHATS_DIR, exist_ok=True)
            with open(_proj_path, "w", encoding="utf-8") as _f:
                json.dump(st.session_state.projects, _f, ensure_ascii=False, indent=2)

        st.markdown("## 🗂️ Projects")
        st.caption("Organise your work into projects. Group related chats under one roof.")
        st.markdown("---")

        # ── Create new project ────────────────────────────────────────────
        with st.expander("➕ New Project", expanded=len(st.session_state.projects) == 0):
            _pname = st.text_input("Project name", key="proj_new_name", placeholder="e.g. Resume Builder")
            _pdesc = st.text_area("Description (optional)", key="proj_new_desc", height=68,
                                  placeholder="What is this project about?")
            if st.button("✅ Create Project", key="proj_create_btn", type="primary"):
                if _pname.strip():
                    st.session_state.projects.append({
                        "id": _new_session_id(),
                        "name": _pname.strip(),
                        "description": _pdesc.strip(),
                        "created_at": datetime.now().strftime("%Y-%m-%d"),
                        "chats": [],
                    })
                    _save_projects()
                    st.success(f"Project **{_pname.strip()}** created!")
                    st.rerun()
                else:
                    st.warning("Enter a project name.")

        st.markdown("---")

        if not st.session_state.projects:
            st.markdown("""
            <div style='text-align:center;padding:3rem 1rem;color:#aaa'>
              <div style='font-size:3.5rem'>🗂️</div>
              <div style='font-size:1.2rem;font-weight:600;margin-top:.5rem'>No projects yet</div>
              <div style='font-size:.9rem'>Create your first project above</div>
            </div>""", unsafe_allow_html=True)
        else:
            # ── Search projects ───────────────────────────────────────────
            _psq = st.text_input("🔍 Search projects", placeholder="Search…",
                                 key="proj_search", label_visibility="collapsed")
            _projs = st.session_state.projects
            if _psq.strip():
                _projs = [p for p in _projs if _psq.lower() in p["name"].lower()
                          or _psq.lower() in p.get("description","").lower()]

            _uname_proj = st.session_state.auth_username or "guest"
            _all_sids = {s["id"]: s["title"] for s in _load_sessions(_uname_proj)}

            for proj in _projs:
                with st.container(border=True):
                    pc1, pc2 = st.columns([5, 1])
                    with pc1:
                        st.markdown(f"### 📁 {proj['name']}")
                        if proj.get("description"):
                            st.caption(proj["description"])
                        st.caption(f"📅 Created {proj.get('created_at','')}")
                    with pc2:
                        if st.button("🗑️", key=f"proj_del_{proj['id']}", help="Delete project"):
                            st.session_state.projects = [
                                p for p in st.session_state.projects if p["id"] != proj["id"]]
                            _save_projects()
                            st.rerun()

                    # ── Attach a chat to this project ─────────────────────
                    with st.expander(f"Chats in this project ({len(proj.get('chats',[]))})"):
                        _linked = proj.get("chats", [])
                        for cid in _linked:
                            cc1, cc2 = st.columns([5, 1])
                            cc1.markdown(f"💬 {_all_sids.get(cid, cid)}")
                            if cc2.button("✖", key=f"proj_rm_{proj['id']}_{cid}", help="Remove"):
                                proj["chats"] = [c for c in proj["chats"] if c != cid]
                                _save_projects()
                                st.rerun()

                        # add a chat
                        _unlinked = [(sid, title) for sid, title in _all_sids.items()
                                     if sid not in _linked]
                        if _unlinked:
                            _pick = st.selectbox("Link a chat", ["— select —"] + [t for _, t in _unlinked],
                                                 key=f"proj_pick_{proj['id']}")
                            if st.button("➕ Link", key=f"proj_link_{proj['id']}"):
                                if _pick != "— select —":
                                    _picked_id = next(sid for sid, t in _unlinked if t == _pick)
                                    proj["chats"].append(_picked_id)
                                    _save_projects()
                                    st.rerun()
                        else:
                            st.caption("All chats already linked.")

                    # ── Rename project ────────────────────────────────────
                    with st.expander("✏️ Rename / Edit"):
                        _new_pname = st.text_input("Name", value=proj["name"], key=f"proj_rn_{proj['id']}")
                        _new_pdesc = st.text_area("Description", value=proj.get("description",""),
                                                  key=f"proj_rd_{proj['id']}", height=60)
                        if st.button("💾 Save", key=f"proj_save_{proj['id']}", type="primary"):
                            proj["name"] = _new_pname.strip() or proj["name"]
                            proj["description"] = _new_pdesc.strip()
                            _save_projects()
                            st.success("Updated!")
                            st.rerun()
        return

    # ── More page ──────────────────────────────────────────────────────────────
    if st.session_state.current_page == "more":
        st.markdown("## ⋯ More")
        st.markdown("---")
        st.markdown("""
        <div style='text-align:center; padding: 4rem 2rem;'>
          <div style='font-size:4rem'>🚀</div>
          <h2 style='color:#1f77b4; margin-top:1rem'>Coming Soon</h2>
          <p style='color:#666; font-size:1.05rem; max-width:400px; margin:1rem auto;'>
            We're building more powerful features for you — Apps, Codex, Collaboration tools and more.
            Stay tuned for updates!
          </p>
          <div style='margin-top:2rem; display:flex; gap:1rem; justify-content:center; flex-wrap:wrap;'>
            <span style='background:#f0f4ff;border:1px solid #c7d7fc;border-radius:20px;
                padding:.4rem 1rem;font-size:.9rem;color:#1f77b4'>📱 Apps</span>
            <span style='background:#f0f4ff;border:1px solid #c7d7fc;border-radius:20px;
                padding:.4rem 1rem;font-size:.9rem;color:#1f77b4'>⌨️ Codex</span>
            <span style='background:#f0f4ff;border:1px solid #c7d7fc;border-radius:20px;
                padding:.4rem 1rem;font-size:.9rem;color:#1f77b4'>🤝 Collaboration</span>
            <span style='background:#f0f4ff;border:1px solid #c7d7fc;border-radius:20px;
                padding:.4rem 1rem;font-size:.9rem;color:#1f77b4'>📊 Analytics</span>
          </div>
        </div>
        """, unsafe_allow_html=True)
        return

    # ── Profile page ──────────────────────────────────────────────────────────
    if st.session_state.current_page == "profile":
        _profile_photo = st.session_state.get("profile_photo", "") or st.session_state.get("profile", {}).get("profile_photo", "")
        if not _profile_photo:
            _profile = get_profile_api(st.session_state.api_url, st.session_state.auth_token)
            if "error" not in _profile:
                _profile_photo = _profile.get("profile_photo", "")
        _uname = st.session_state.auth_username or "U"
        _initials = "".join(p[0].upper() for p in _uname.split()[:2]) or _uname[:2].upper()
        _avatar_html = (f"""<img src="{_profile_photo}" style="width:64px;height:64px;object-fit:cover;border-radius:50%;border:1px solid #e0e0e0;" />"""
                        if _profile_photo
                        else f"""<div style="width:64px;height:64px;border-radius:50%;background:#1f77b4;color:white;font-size:26px;font-weight:700;display:inline-flex;align-items:center;justify-content:center;border:1px solid #e0e0e0;vertical-align:middle;">{_initials}</div>""")
        st.markdown(f"## {_avatar_html} User Profile", unsafe_allow_html=True)
        if st.button("← Back to Chat", key="back_to_chat"):
            st.session_state.current_page = "chat"
            st.rerun()
        st.markdown("---")

        profile = get_profile_api(st.session_state.api_url, st.session_state.auth_token)
        if "error" in profile:
            st.error(profile["error"])
        else:
            col1, col2 = st.columns(2)
            with col1:
                st.markdown("### Account Info")
                st.text_input("Username", value=profile.get("username", ""), disabled=True, key="prof_username")
                st.text_input("Role", value=profile.get("role", ""), disabled=True, key="prof_role")
                joined = profile.get("created_at", "")[:10] if profile.get("created_at") else "—"
                st.text_input("Member Since", value=joined, disabled=True, key="prof_joined")

            with col2:
                st.markdown("### Edit Profile")
                new_display = st.text_input("Display Name", value=profile.get("display_name", ""), key="prof_display")
                new_email = st.text_input("Email Address", value=profile.get("email", ""), key="prof_email")

                st.markdown("**Profile Photo**")
                _photo_init = st.session_state.get("profile_photo", profile.get("profile_photo", ""))
                new_photo_url = st.text_input("Photo URL", value=_photo_init, key="prof_photo_url",
                                              placeholder="https://example.com/photo.jpg")
                up1, up2 = st.columns([3, 2])
                with up1:
                    _photo_file = st.file_uploader("Or upload photo", type=["png","jpg","jpeg","svg"],
                                                   key="prof_photo_file")
                if _photo_file:
                    _photo_b64 = base64.b64encode(_photo_file.read()).decode()
                    _photo_mime = _photo_file.type
                    new_photo_url = f"data:{_photo_mime};base64,{_photo_b64}"
                    with up2:
                        st.markdown("""<div style="font-size:0.8rem;color:#666;margin-bottom:2px;">Preview</div>""", unsafe_allow_html=True)
                        st.markdown(f"""<img src="{new_photo_url}" style="width:90px;height:90px;object-fit:cover;border-radius:50%;border:2px solid #1f77b4;" />""", unsafe_allow_html=True)
                        st.caption("Click 💾 Save Info to confirm")

                st.markdown("**Company / Organisation**")
                _co_init = st.session_state.get("company_name", profile.get("company_name", ""))
                new_company = st.text_input("Company Name", value=_co_init, key="prof_company",
                                            placeholder="e.g. Acme Corp")
                _logo_url_init = st.session_state.get("company_logo", profile.get("company_logo", ""))
                new_logo_url = st.text_input("Logo URL", value=_logo_url_init, key="prof_logo_url",
                                             placeholder="https://example.com/logo.png")
                _logo_file = st.file_uploader("Or upload logo image", type=["png","jpg","jpeg","svg"],
                                              key="prof_logo_file")
                if _logo_file:
                    _logo_b64 = base64.b64encode(_logo_file.read()).decode()
                    _logo_mime = _logo_file.type
                    new_logo_url = f"data:{_logo_mime};base64,{_logo_b64}"
                _website_init = st.session_state.get("company_website", profile.get("company_website", ""))
                new_website = st.text_input("Company URL", value=_website_init, key="prof_website",
                                            placeholder="https://www.example.com")
                _mobile_init = profile.get("mobile", "")

                # Auto-detect country once per session (returns full label string)
                if "detected_dial" not in st.session_state:
                    st.session_state.detected_dial = _detect_country_dial()

                # Split saved number back into dial+local for display
                _default_label = st.session_state.detected_dial
                _saved_local = ""
                if _mobile_init:
                    # Try to match stored number against known dial codes (longest first)
                    for lbl in sorted(_DIAL_LABELS, key=lambda l: len(_DIAL_MAP[l]), reverse=True):
                        _code = _DIAL_MAP[lbl]
                        if _mobile_init.startswith(_code + " ") or _mobile_init.startswith(_code + "-"):
                            _default_label = lbl
                            _saved_local = _mobile_init[len(_code):].strip().lstrip("-")
                            break
                    else:
                        _saved_local = "".join(c for c in _mobile_init if c.isdigit() or c in "- ")

                # Ensure default label exists in list (fallback to US if not)
                if _default_label not in _DIAL_LABELS:
                    _default_label = _DIAL_LABELS[0]
                _default_idx = _DIAL_LABELS.index(_default_label)

                st.markdown("**Mobile Number**")
                mob_c1, mob_c2 = st.columns([2, 3])
                with mob_c1:
                    selected_label = st.selectbox(
                        "Country", options=_DIAL_LABELS,
                        index=_default_idx, key="prof_dial_code",
                        label_visibility="collapsed"
                    )
                    chosen_dial = _DIAL_MAP[selected_label]
                with mob_c2:
                    local_number = st.text_input(
                        "Local number", value=_saved_local,
                        key="prof_mobile_local", placeholder="e.g. 1712 345 678",
                        label_visibility="collapsed"
                    )

                new_mobile = _format_phone(chosen_dial, local_number)
                if new_mobile:
                    st.caption(f"📱 Full number: `{new_mobile}`")
                if st.button("💾 Save Info", type="primary", use_container_width=True, key="save_info_btn"):
                    # Store company info in session state so header updates immediately
                    st.session_state.company_name = new_company.strip()
                    st.session_state.company_logo = new_logo_url.strip()
                    st.session_state.company_website = new_website.strip()
                    st.session_state.profile_photo = new_photo_url.strip()
                    res = update_profile_api(
                        st.session_state.api_url,
                        st.session_state.auth_token,
                        {
                            "display_name": new_display.strip() or None,
                            "email": new_email.strip() or None,
                            "phone": new_mobile or None,
                        }
                    )
                    if "error" in res:
                        st.error(res["error"])
                    else:
                        st.success("✅ Profile updated successfully!")

            st.markdown("---")
            st.markdown("### 🔒 Change Password")
            pw_col1, pw_col2 = st.columns(2)
            with pw_col1:
                cur_pw = st.text_input("Current Password", type="password", key="cur_pw")
                new_pw = st.text_input("New Password", type="password", key="new_pw")
            with pw_col2:
                confirm_pw = st.text_input("Confirm New Password", type="password", key="confirm_pw")
                st.markdown("<br>", unsafe_allow_html=True)
                if st.button("🔑 Change Password", type="primary", use_container_width=True, key="change_pw_btn"):
                    if not (cur_pw and new_pw and confirm_pw):
                        st.warning("Fill in all password fields.")
                    elif new_pw != confirm_pw:
                        st.error("New passwords do not match.")
                    else:
                        res = update_profile_api(
                            st.session_state.api_url,
                            st.session_state.auth_token,
                            {"current_password": cur_pw, "new_password": new_pw}
                        )
                        if "error" in res:
                            st.error(res["error"])
                        else:
                            st.success("✅ Password changed successfully!")
        return

    # ── Admin panel page ──────────────────────────────────────────────────────
    if st.session_state.current_page == "admin" and st.session_state.get("auth_role") == "admin":
        st.markdown("## ⚙️ Admin Panel")
        if st.button("← Back to Chat", key="admin_back"):
            st.session_state.current_page = "chat"
            st.rerun()
        st.markdown("---")
        st.markdown("### User Accounts")
        _users_res = admin_list_users_api(st.session_state.api_url, st.session_state.auth_token)
        if "error" in _users_res:
            st.error(_users_res["error"])
        else:
            _users = _users_res.get("users", [])
            for _u in _users:
                _c1, _c2, _c3, _c4 = st.columns([3, 2, 1, 1])
                with _c1:
                    st.markdown(f"**{_u['username']}**")
                    st.caption(f"{_u.get('display_name', '')} | {_u.get('company_name', '')}")
                with _c2:
                    _role_color = {"admin": "🟥", "user": "🟦"}.get(_u['role'], "⬜")
                    st.markdown(f"{_role_color} {_u['role']}")
                    st.caption(_u.get('created_at', '')[:10])
                with _c3:
                    _new_role = st.selectbox("Role", ["user", "admin"], index=["user", "admin"].index(_u['role']),
                                             key=f"role_{_u['username']}", label_visibility="collapsed")
                    if _new_role != _u['role']:
                        if st.button("Update", key=f"upd_role_{_u['username']}"):
                            _res = admin_set_role_api(st.session_state.api_url, st.session_state.auth_token, _u['username'], _new_role)
                            if "error" in _res:
                                st.error(_res["error"])
                            else:
                                st.success("Role updated")
                                st.rerun()
                with _c4:
                    if _u['username'] != st.session_state.auth_username:
                        if st.button("🗑", key=f"del_user_{_u['username']}"):
                            st.session_state._pending_delete_user = _u['username']
                            st.rerun()
        st.markdown("---")

        # ── User deletion confirmation dialog ─────────────────────────────
        if st.session_state.get("_pending_delete_user"):
            @st.dialog("Confirm deletion")
            def _confirm_delete():
                _target = st.session_state._pending_delete_user
                st.write(f"Are you sure you want to permanently delete user **{_target}**?")
                _c1, _c2 = st.columns(2)
                with _c1:
                    if st.button("Yes, delete", type="primary", key="confirm_delete_user"):
                        _res = admin_delete_user_api(st.session_state.api_url, st.session_state.auth_token, _target)
                        if "error" in _res:
                            st.error(_res["error"])
                        else:
                            st.session_state.pop("_pending_delete_user", None)
                            st.success("User deleted")
                            st.rerun()
                with _c2:
                    if st.button("Cancel", key="cancel_delete_user"):
                        st.session_state.pop("_pending_delete_user", None)
                        st.rerun()
            _confirm_delete()
        st.markdown("---")
        st.markdown("### Health Monitoring")
        _health_res = admin_health_status_api(st.session_state.api_url, st.session_state.auth_token)
        if "error" in _health_res:
            st.error(_health_res["error"])
        else:
            if not _health_res.get("monitoring_enabled"):
                st.info("Monitoring is disabled. Set MONITORING_ENABLED=true in .env to enable it.")
            else:
                _healthy = _health_res.get("healthy", False)
                st.markdown(f"**Overall:** {'✅ Healthy' if _healthy else '❌ Issues detected'}")
                _latest = _health_res.get("latest", {})
                if _latest:
                    for _comp, _info in _latest.items():
                        _status = _info.get("status", "unknown")
                        _emoji = {"ok": "🟢", "warning": "🟡", "error": "🔴"}.get(_status, "⚪")
                        _msg = _info.get("message", "")
                        _rt = _info.get("response_time_ms", 0)
                        st.caption(f"{_emoji} {_comp}: {_status} — {_msg} ({_rt:.1f} ms)")
                else:
                    st.caption("No health checks yet. They run every few seconds.")
        st.markdown("---")
        st.markdown("### Embeddable Widget")
        _widget_res = admin_widget_config_api(st.session_state.api_url, st.session_state.auth_token)
        if "error" in _widget_res:
            st.error(_widget_res["error"])
        else:
            if not _widget_res.get("enabled"):
                st.info(_widget_res.get("message", "Widget is disabled."))
            else:
                st.caption("Copy this snippet into any website to embed the chat widget.")
                st.code(_widget_res.get("embed_code", ""), language="html")
                if st.button("📋 Copy embed code", key="copy_widget_code"):
                    st.toast("Embed code copied to clipboard (use Ctrl+C on the code block)")
        st.markdown("---")
        st.markdown("### User Feedback")
        _feedback_res = admin_list_feedback_api(st.session_state.api_url, st.session_state.auth_token)
        if "error" in _feedback_res:
            st.error(_feedback_res["error"])
        else:
            _feedback_items = _feedback_res.get("feedback", [])
            if not _feedback_items:
                st.caption("No feedback yet.")
            else:
                st.caption(f"Total feedback entries: {len(_feedback_items)}")
                # Sort by newest first
                for _fb in reversed(_feedback_items):
                    _rating = _fb.get("rating", "")
                    _icon = "👍" if _rating == "up" else "👎" if _rating == "down" else "❓"
                    with st.expander(f"{_icon} {_fb.get('username', 'unknown')} — {_fb.get('timestamp', '')[:19]}"):
                        st.markdown(f"**Mode:** {_fb.get('mode', 'document')}")
                        st.markdown(f"**Question:** {_fb.get('query', '')}")
                        st.markdown(f"**Answer:** {_fb.get('response', '')}")
                        if _fb.get("comment"):
                            st.markdown(f"**Comment:** {_fb.get('comment')}")
        st.markdown("### Tenant Migration")
        st.markdown("If you previously used the shared ChromaDB collection, run this once to migrate data into isolated per-user collections.")
        if st.button("Run migration", key="run_migrate"):
            _res = requests.post(f"{st.session_state.api_url}/api/admin/migrate-to-tenant-collections",
                                 headers={"Authorization": f"Bearer {st.session_state.auth_token}"}, timeout=120)
            if _res.status_code == 200:
                st.success(f"Migrated {_res.json().get('migrated_chunks', 0)} chunks.")
            else:
                st.error(_res.json().get("detail", "Migration failed."))
        return

    # ── Document preview panel ────────────────────────────────────────────
    if st.session_state.get("doc_preview_name"):
        _pname    = st.session_state["doc_preview_name"]
        _pcontent = st.session_state.get("doc_preview_content", "")
        _ptype    = st.session_state.get("doc_preview_type", "text")
        _ppages   = st.session_state.get("doc_preview_pages", "")

        with st.expander(f"📄 Previewing: {_pname}", expanded=True):
            hdr_c1, hdr_c2 = st.columns([9, 1])
            with hdr_c1:
                _type_icon = {"pdf": "📑 PDF", "docx": "📝 Word", "text": "📃 Text"}.get(_ptype, "📄")
                _pg_info = f" — {_ppages} pages" if _ppages else ""
                st.caption(f"{_type_icon}{_pg_info} — extracted text preview")
            with hdr_c2:
                if st.button("✖ Close", key="close_doc_preview"):
                    st.session_state["doc_preview_name"] = None
                    st.session_state["doc_preview_content"] = ""
                    st.session_state["doc_preview_pages"] = ""
                    st.rerun()

            # Render as styled scrollable HTML — preserves whitespace, wraps properly
            import html as _html
            _escaped = _html.escape(_pcontent)
            # Turn page separators into visible dividers
            _escaped = _escaped.replace("── Page ", "<hr style='margin:8px 0;border-color:#ccc'><strong>── Page ")
            _escaped = _escaped.replace(" ──\n", " ──</strong><br>")
            st.markdown(f"""
            <div style="
                background:#fafafa;
                border:1px solid #e0e0e0;
                border-radius:6px;
                padding:16px 20px;
                height:500px;
                overflow-y:auto;
                font-family:'Segoe UI',Arial,sans-serif;
                font-size:0.92rem;
                line-height:1.7;
                color:#222;
                white-space:pre-wrap;
                word-break:break-word;
            ">{_escaped}</div>""", unsafe_allow_html=True)
        st.markdown("---")

    # ── Document chunk viewer panel ──────────────────────────────────────
    if st.session_state.get("doc_chunks_name"):
        _cname = st.session_state["doc_chunks_name"]
        if st.session_state.get("doc_chunks_data") is None:
            with st.spinner(f"Loading chunks for {_cname}..."):
                _chunks_res = get_document_chunks_api(
                    st.session_state.api_url, st.session_state.auth_token, _cname
                )
            st.session_state["doc_chunks_data"] = _chunks_res
            st.rerun()
        else:
            _chunks_res = st.session_state["doc_chunks_data"]
            if "error" in _chunks_res:
                st.error(_chunks_res["error"])
            else:
                _chunks = _chunks_res.get("chunks", [])
                _count = _chunks_res.get("count", 0)
                with st.expander(f"🧩 Chunks: {_cname} ({_count} chunk(s))", expanded=True):
                    hdr_c1, hdr_c2 = st.columns([9, 1])
                    with hdr_c1:
                        st.caption("Indexed chunks stored in ChromaDB")
                    with hdr_c2:
                        if st.button("✖ Close", key="close_doc_chunks"):
                            st.session_state["doc_chunks_name"] = None
                            st.session_state["doc_chunks_data"] = None
                            st.rerun()
                    if not _chunks:
                        st.info("No indexed chunks found for this document. Run Sync Index first.")
                    else:
                        for _idx, _chunk in enumerate(_chunks, 1):
                            _text = _chunk.get("text", "")
                            _emb_len = _chunk.get("embedding_length", 0)
                            with st.container(border=True):
                                st.markdown(f"**Chunk {_idx}** · `{_chunk.get('chunk_id', '')}` · embedding dim: `{_emb_len}`")
                                st.text(_text[:2000] + ("…" if len(_text) > 2000 else ""))
            st.markdown("---")

    @st.fragment()
    def _render_chat_area():
        """Render the chat session and input. Runs as a Streamlit fragment so
        sending a message only rerenders this area, not the whole sidebar."""
        _uname2 = st.session_state.auth_username or "guest"
        if st.session_state.active_session_id is None:
            st.session_state.active_session_id = _new_session_id()
        _active_sess = _load_session(_uname2, st.session_state.active_session_id)
        _sess_title = _active_sess.get("title", "New Chat")
        if _sess_title != "New Chat":
            st.markdown(f"<h3 style='margin-top:0;color:#555;font-size:1.1rem'>{_sess_title}</h3>",
                        unsafe_allow_html=True)

        # Auto-scroll inside the fragment after a new message is added
        if st.session_state.get("scroll_to_bottom"):
            st.session_state.scroll_to_bottom = False
            components.html(
                """<script>
                (function() {
                    function findAndScroll() {
                        var parentDoc = window.parent.document;
                        var all = parentDoc.querySelectorAll('*');
                        for (var i = 0; i < all.length; i++) {
                            var el = all[i];
                            var style = window.parent.getComputedStyle(el);
                            var overflow = style.overflow + style.overflowY;
                            if ((overflow.indexOf('scroll') !== -1 || overflow.indexOf('auto') !== -1)
                                    && el.scrollHeight > el.clientHeight) {
                                el.scrollTop = el.scrollHeight;
                            }
                        }
                        window.parent.scrollTo(0, window.parent.document.body.scrollHeight);
                    }
                    setTimeout(findAndScroll, 300);
                    setTimeout(findAndScroll, 800);
                })();
                </script>""",
                height=0,
            )

        # ── Conversation history (renders at top, grows downward) ────────────────
        _msgs = _active_sess.get("messages", [])
        _sid = st.session_state.active_session_id
        _uname_chat = st.session_state.auth_username or "guest"
        if _msgs:
            for idx, entry in enumerate(_msgs):
                mode_label = "📄 From My Documents" if entry.get("mode") == "document" else "🤖 Ask AI Freely"
                _menu_key = f"msg_menu_{_sid}_{idx}"
                _edit_key = f"msg_edit_{_sid}_{idx}"

                # ── User bubble row: question + three-dot menu ────────────────
                u_col, dot_col = st.columns([11, 1])
                with u_col:
                    with st.chat_message("user"):
                        st.caption(mode_label)
                        if st.session_state.get(_edit_key, False):
                            _new_q = st.text_area(
                                "Edit question", value=entry["question"],
                                key=f"edit_input_{_sid}_{idx}",
                                label_visibility="collapsed", height=80
                            )
                            sv_c, cn_c = st.columns(2)
                            with sv_c:
                                if st.button("💾 Save", key=f"edit_save_{_sid}_{idx}", use_container_width=True):
                                    _active_sess["messages"][idx]["question"] = _new_q.strip()
                                    _save_session(_uname_chat, _active_sess)
                                    st.session_state[_edit_key] = False
                                    st.session_state[_menu_key] = False
                                    st.rerun(scope="fragment")
                            with cn_c:
                                if st.button("✖ Cancel", key=f"edit_cancel_{_sid}_{idx}", use_container_width=True):
                                    st.session_state[_edit_key] = False
                                    st.session_state[_menu_key] = False
                                    st.rerun(scope="fragment")
                        else:
                            st.markdown(entry["question"])

                with dot_col:
                    st.markdown("<div style='padding-top:18px'>", unsafe_allow_html=True)
                    if st.button("⋯", key=f"dot_{_sid}_{idx}", help="Edit / Delete"):
                        st.session_state[_menu_key] = not st.session_state.get(_menu_key, False)
                        st.session_state[_edit_key] = False
                    st.markdown("</div>", unsafe_allow_html=True)

                # ── Pop-up action row ─────────────────────────────────────────
                if st.session_state.get(_menu_key, False):
                    ma, mb, mc = st.columns([1, 1, 9])
                    with ma:
                        if st.button("✏️ Edit", key=f"menu_edit_{_sid}_{idx}", use_container_width=True):
                            st.session_state[_edit_key] = True
                            st.session_state[_menu_key] = False
                            st.rerun(scope="fragment")
                    with mb:
                        if st.button("🗑️ Delete", key=f"menu_del_{_sid}_{idx}", use_container_width=True):
                            _active_sess["messages"].pop(idx)
                            _save_session(_uname_chat, _active_sess)
                            st.session_state[_menu_key] = False
                            st.rerun(scope="fragment")

                # ── Assistant bubble ──────────────────────────────────────────
                with st.chat_message("assistant"):
                    st.markdown(entry["answer"])
                    fname = f"kb_answer_{idx+1}"
                    q, a = entry["question"], entry["answer"]
                    ec1, ec2, ec3, ec4 = st.columns(4)
                    with ec1:
                        st.download_button(
                            "⬇ TXT", data=_export_txt(q, a),
                            file_name=f"{fname}.txt", mime="text/plain",
                            key=f"exp_txt_{_sid}_{idx}", use_container_width=True
                        )
                    with ec2:
                        st.download_button(
                            "⬇ MD", data=_export_md(q, a),
                            file_name=f"{fname}.md", mime="text/markdown",
                            key=f"exp_md_{_sid}_{idx}", use_container_width=True
                        )
                    with ec3:
                        st.download_button(
                            "⬇ DOCX", data=_export_docx(q, a),
                            file_name=f"{fname}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"exp_docx_{_sid}_{idx}", use_container_width=True
                        )
                    with ec4:
                        st.download_button(
                            "⬇ PDF", data=_export_pdf(q, a),
                            file_name=f"{fname}.pdf", mime="application/pdf",
                            key=f"exp_pdf_{_sid}_{idx}", use_container_width=True
                        )

                    # ── Feedback row ────────────────────────────────────────────
                    fb = entry.get("feedback_rating")
                    fb_col1, fb_col2, fb_col3 = st.columns([1, 1, 6])
                    with fb_col1:
                        disabled_up = fb is not None
                        if st.button("👍", key=f"fb_up_{_sid}_{idx}", help="Helpful answer", disabled=disabled_up, use_container_width=True):
                            submit_feedback(
                                st.session_state.api_url, st.session_state.auth_token,
                                q, a, entry.get("mode", "document"), "up",
                                session_id=_sid, message_id=str(idx)
                            )
                            _active_sess["messages"][idx]["feedback_rating"] = "up"
                            _save_session(_uname_chat, _active_sess)
                            st.toast("Thanks for your feedback!")
                            st.rerun(scope="fragment")
                    with fb_col2:
                        disabled_down = fb is not None
                        if st.button("👎", key=f"fb_down_{_sid}_{idx}", help="Not helpful", disabled=disabled_down, use_container_width=True):
                            st.session_state[f"fb_down_open_{_sid}_{idx}"] = True
                            st.rerun(scope="fragment")
                    if st.session_state.get(f"fb_down_open_{_sid}_{idx}", False):
                        with fb_col3:
                            comment = st.text_input(
                                "What was wrong? (optional)",
                                key=f"fb_comment_{_sid}_{idx}",
                                placeholder="e.g., wrong answer, missing citation...",
                                label_visibility="collapsed"
                            )
                            _fb_col3_a, _fb_col3_b = st.columns([1, 1])
                            with _fb_col3_a:
                                if st.button("Submit feedback", key=f"fb_submit_{_sid}_{idx}", type="primary", use_container_width=True):
                                    submit_feedback(
                                        st.session_state.api_url, st.session_state.auth_token,
                                        q, a, entry.get("mode", "document"), "down", comment=comment,
                                        session_id=_sid, message_id=str(idx)
                                    )
                                    _active_sess["messages"][idx]["feedback_rating"] = "down"
                                    _save_session(_uname_chat, _active_sess)
                                    st.session_state[f"fb_down_open_{_sid}_{idx}"] = False
                                    st.toast("Thanks for your feedback!")
                                    st.rerun(scope="fragment")
                            with _fb_col3_b:
                                if st.button("Cancel", key=f"fb_cancel_{_sid}_{idx}", use_container_width=True):
                                    st.session_state[f"fb_down_open_{_sid}_{idx}"] = False
                                    st.rerun(scope="fragment")

        # ── Streaming placeholder (filled during active search) ──────────────────
        stream_placeholder = st.empty()

        # ── Input area pinned at the bottom ──────────────────────────────────────
        st.markdown("---")
        st.markdown("### Ask a Question")

        user_query = st.text_area(
            "Enter your question:",
            value=st.session_state.get("query_input", ""),
            placeholder="e.g., What technologies are used in this pilot stack?",
            height=80,
            key="query_input",
            label_visibility="visible"
        )

        mode = st.radio(
            "Mode:",
            options=["document", "assistant"],
            format_func=lambda m: "📄 From My Documents" if m == "document" else "🤖 Ask AI Freely",
            horizontal=True,
            key="query_mode",
        )

        col1, col2 = st.columns([2, 1])
        with col1:
            st.markdown('<style>div[data-testid="stButton"] > button[kind="primary"] { background-color: #1f77b4 !important; color: white !important; }</style>', unsafe_allow_html=True)
            search_button = st.button("🔍 Search", type="primary", use_container_width=True, key="search_button", disabled=not st.session_state.backend_connected)
        with col2:
            clear_button = st.button("🗑️ Clear", use_container_width=True, key="clear_button")

        if clear_button:
            _active_sess["messages"] = []
            _active_sess["title"] = "New Chat"
            st.session_state.query_input = ""
            _save_session(_uname2, _active_sess)
            st.rerun(scope="fragment")

        if search_button:
            logger.info("Search button clicked")
            if not user_query.strip():
                st.warning("Please enter a question.")
                return

            live_status = check_backend_connection(st.session_state.api_url)
            st.session_state.backend_connected = live_status["connected"]
            st.session_state.connection_status = live_status
            if not live_status["connected"]:
                st.error("Backend is not connected. Please wait for the backend to start.")
                st.rerun(scope="fragment")
                return

            try:
                logger.info("Streaming query started")
                mode_label = "📄 From My Documents" if mode == "document" else "🤖 Ask AI Freely"
                history_payload = []
                for e in _active_sess.get("messages", []):
                    history_payload.append({"role": "user", "content": e["question"]})
                    history_payload.append({"role": "assistant", "content": e["answer"]})
                with stream_placeholder.container():
                    with st.chat_message("user"):
                        st.caption(mode_label)
                        st.markdown(user_query)
                    with st.chat_message("assistant"):
                        answer_placeholder = st.empty()
                        accumulated = ""
                        last_render_len = 0
                        for token in stream_query(user_query, st.session_state.api_url, mode=mode, history=history_payload, token=st.session_state.auth_token):
                            accumulated += token
                            if len(accumulated) - last_render_len >= 20:
                                answer_placeholder.markdown(accumulated + "▌")
                                last_render_len = len(accumulated)
                        answer_placeholder.markdown(accumulated)
                new_entry = {"question": user_query, "answer": _sanitize_answer(accumulated), "mode": mode, "feedback_rating": None}
                _active_sess["messages"].append(new_entry)
                if len(_active_sess["messages"]) == 1:
                    _active_sess["title"] = _auto_title(user_query)
                _save_session(_uname2, _active_sess)
                st.session_state.query_input = ""
                st.session_state.input_counter += 1
                st.session_state.scroll_to_bottom = True
                logger.info("Streaming query completed")
                st.rerun(scope="fragment")
            except requests.exceptions.HTTPError as e:
                logger.error(f"Streaming query failed (HTTP {e.response.status_code}): {e}")
                if e.response.status_code == 401:
                    st.warning("Session expired. Please log in again.")
                    st.session_state.auth_token = None
                    st.session_state.auth_refresh_token = None
                    st.session_state.auth_username = None
                    st.session_state.auth_role = None
                    st.query_params.pop("_t", None)
                    st.rerun()
                else:
                    st.error(f"⚠️ Backend returned an error: {e.response.status_code}. Please check the server logs.")
                return
            except requests.exceptions.RequestException as e:
                logger.error(f"Streaming query failed (network): {e}")
                st.error(
                    "⚠️ **Could not reach the backend.**\n\n"
                    "Please check that the server and Ollama are both running, then try again."
                )
            return

        # Example queries
        st.markdown("### Example Questions")
        st.markdown('<style>div[data-testid="stVerticalBlock"] > div[data-testid="stVerticalBlock"] { border-top: none !important; }</style>', unsafe_allow_html=True)
        example_queries = [
            "What technologies are used in this pilot stack?",
            "When was the project kickoff date?",
            "What is the purpose of this knowledge base?",
            "How does the RAG system work?"
        ]
        col1, col2 = st.columns([1, 1], gap="small")
        for i, query in enumerate(example_queries):
            if i % 2 == 0:
                with col1:
                    st.markdown('<style>div[data-testid="column"] { text-align: left !important; }</style>', unsafe_allow_html=True)
                    st.button(query, key=f"example_{i}", type="secondary", use_container_width=True, on_click=set_example_query, args=(query,))
            else:
                with col2:
                    st.markdown('<style>div[data-testid="column"] { text-align: left !important; }</style>', unsafe_allow_html=True)
                    st.button(query, key=f"example_{i}", type="secondary", use_container_width=True, on_click=set_example_query, args=(query,))

    _render_chat_area()

if __name__ == "__main__":
    main()
